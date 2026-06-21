from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import Settings, get_settings
from app.db.base import utc_now
from app.models import (
    AUDPlan,
    AUDSectionDraft,
    EvidenceItem,
    ExtractedContent,
    GeneratedDocument,
    OpenPoint,
    Project,
)
from app.services.file_storage import (
    LocalStorageService,
    StorageService,
    get_file_storage,
)
from app.services.docx_table_renderer import (
    DOCXTableRenderer,
    NormalizedTable,
    TableNormalizer,
)
from app.services.template_resolver import TemplateResolver
from app.services.template_population import (
    PopulatedDocumentModel,
    PopulatedImage,
    PopulatedOpenPoint,
    PopulatedSection,
    PopulatedTable,
    SectionPopulationInput,
    TemplatePopulationService,
    clean_template_text,
)

DOCUMENT_TYPE = "aud_docx"
PLACEHOLDER_LINE = "<Content not available in provided source material>"
GENERATION_NOTE = (
    "Generated draft for Oracle internal review. Senior consultant review required "
    "before customer sharing."
)
SKIP_PLANNED_SECTION_TITLES = {
    "cover page",
    "documents referred",
    "document version history",
    "table of contents",
    "purpose and scope",
    "open points",
}
MAX_SECTION_PARAGRAPHS = 4
MAX_SECTION_CHARACTERS = 1800
MAX_ENTERPRISE_SECTION_PARAGRAPHS = 20
MAX_ENTERPRISE_SECTION_CHARACTERS = 8000
MAX_IMAGES_PER_SECTION = 3
SECTION_HEADING_FONT_NAME = "Oracle Sans"
SECTION_HEADING_FONT_SIZE_PT = 16
SECTION_HEADING_COLOR = RGBColor(31, 78, 121)
SECTION_HEADING_SPACE_BEFORE_PT = 18
SECTION_HEADING_SPACE_AFTER_PT = 8
CONTENT_SUBHEADING_FONT_SIZE_PT = 14
CONTENT_SUBHEADING_COLOR = RGBColor(31, 78, 121)
STEP_HEADING_FONT_SIZE_PT = 13
STEP_HEADING_COLOR = RGBColor(68, 114, 196)
BODY_SPACE_AFTER_PT = 6
ENTERPRISE_STRUCTURE_TITLE = "Enterprise Structure"
SUPPORTED_DOCX_IMAGE_EXTENSIONS = {
    ".bmp",
    ".gif",
    ".jpeg",
    ".jpg",
    ".png",
    ".tif",
    ".tiff",
}
LOW_VALUE_SLIDE_TITLES = {
    "agenda",
    "agenda only",
    "divider",
    "thank you",
    "thanks",
    "welcome",
}
ACCEPTED_DRAFT_STATUSES = {"accepted", "approved", "reviewed"}
OMITTED_DRAFT_STATUSES = {"excluded", "omitted", "removed"}
TEMPLATE_PLACEHOLDER_PATTERN = re.compile(r"<+[^<>]+>+")
TEMPLATE_BODY_START_TITLES = {
    "document version history",
    "table of contents",
    "table of content",
    "introduction",
    "purpose and scope",
}
logger = logging.getLogger(__name__)
TABLE_NORMALIZER = TableNormalizer()
TABLE_RENDERER = DOCXTableRenderer(TABLE_NORMALIZER)

@dataclass(frozen=True)
class DocxGenerationOptions:
    use_ai_drafts: bool = True
    include_draft_sections: bool = True
    include_images: bool = True
    include_open_points: bool = True


@dataclass(frozen=True)
class SectionRenderContext:
    draft: AUDSectionDraft | None
    draft_json: dict[str, Any]
    paragraphs: list[str]
    tables: list[PopulatedTable]
    was_truncated: bool


@dataclass(frozen=True)
class OpenPointSelection:
    open_points: list[OpenPoint]
    used_fallback: bool


def normalize_title(value: str) -> str:
    return " ".join(value.strip().split()).lower()


def title_tokens(value: str | None) -> set[str]:
    if not value:
        return set()

    stop_words = {
        "a",
        "an",
        "and",
        "for",
        "in",
        "of",
        "on",
        "the",
        "to",
        "with",
    }
    return {
        token
        for token in re.findall(r"[a-z0-9]+", normalize_title(value))
        if token not in stop_words
    }


def titles_resemble(left: str, right: str | None) -> bool:
    left_normalized = normalize_title(left)
    right_normalized = normalize_title(right or "")

    if not left_normalized or not right_normalized:
        return False

    left_tokens = title_tokens(left_normalized)
    right_tokens = title_tokens(right_normalized)

    if left_normalized == right_normalized:
        return True

    if (
        min(len(left_tokens), len(right_tokens)) >= 2
        and (left_normalized in right_normalized or right_normalized in left_normalized)
    ):
        return True

    if not left_tokens or not right_tokens:
        return False

    overlap = len(left_tokens & right_tokens)
    return overlap / max(len(left_tokens), len(right_tokens)) >= 0.6


def is_enterprise_structure_title(value: str | None) -> bool:
    return bool(value and titles_resemble(value, ENTERPRISE_STRUCTURE_TITLE))


def sanitize_filename_part(value: str | None) -> str:
    if not value:
        return "aud"

    sanitized = re.sub(r"[^A-Za-z0-9]+", "-", value.strip()).strip("-").lower()
    return sanitized or "aud"


def parse_plan_json(aud_plan: AUDPlan | None) -> dict[str, Any]:
    if aud_plan is None:
        return {}

    try:
        parsed = json.loads(aud_plan.plan_json)
    except json.JSONDecodeError:
        return {}

    return parsed if isinstance(parsed, dict) else {}


def parse_json_object(value: str | None) -> dict[str, Any]:
    if not value:
        return {}

    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return {}

    return parsed if isinstance(parsed, dict) else {}


def clean_generated_text(value: str | None) -> str:
    return clean_template_text(value)


def dedupe_consecutive_text(items: list[str]) -> list[str]:
    deduped: list[str] = []
    previous_normalized = ""

    for item in items:
        normalized = normalize_title(item)
        if normalized and normalized == previous_normalized:
            continue

        deduped.append(item)
        previous_normalized = normalized

    return deduped


def parse_docx_generation_options(value: str | None = None) -> DocxGenerationOptions:
    payload = parse_json_object(value)
    options = payload.get("options") if payload else {}

    if not isinstance(options, dict):
        options = {}

    return DocxGenerationOptions(
        use_ai_drafts=bool(options.get("use_ai_drafts", True)),
        include_draft_sections=bool(options.get("include_draft_sections", True)),
        include_images=bool(options.get("include_images", True)),
        include_open_points=bool(options.get("include_open_points", True)),
    )


def plan_has_content_sections(plan_payload: dict[str, Any]) -> bool:
    return bool(iter_planned_sections(plan_payload))


def plan_has_enterprise_structure_section(plan_payload: dict[str, Any]) -> bool:
    return any(
        is_enterprise_structure_title(get_section_title(section))
        for section in iter_planned_sections(plan_payload)
    )


def parse_extracted_json(extracted_content: ExtractedContent) -> dict[str, Any]:
    return parse_json_object(extracted_content.json_content)


def latest_aud_plan(session: Session, project_id: str) -> AUDPlan | None:
    statement = (
        select(AUDPlan)
        .where(AUDPlan.project_id == project_id)
        .order_by(AUDPlan.created_at.desc())
    )
    return session.scalars(statement).first()


def list_open_points_for_docx(
    session: Session,
    project_id: str,
    settings: Settings,
) -> OpenPointSelection:
    statement = (
        select(OpenPoint)
        .where(OpenPoint.project_id == project_id)
        .order_by(OpenPoint.created_at.asc())
    )
    open_points = [
        open_point
        for open_point in session.scalars(statement)
        if normalize_title(open_point.status) == "open"
    ]
    enhanced_open_points = [
        open_point
        for open_point in open_points
        if open_point.source_type == "llm_enhanced"
    ]

    if enhanced_open_points:
        return OpenPointSelection(enhanced_open_points, used_fallback=False)

    return OpenPointSelection([], used_fallback=False)


def list_section_drafts(session: Session, project_id: str) -> list[AUDSectionDraft]:
    return list(
        session.scalars(
            select(AUDSectionDraft)
            .where(AUDSectionDraft.project_id == project_id)
            .order_by(AUDSectionDraft.created_at.asc())
        )
    )


def get_evidence_item_by_id(
    session: Session,
    project_id: str,
) -> dict[str, EvidenceItem]:
    evidence_items = session.scalars(
        select(EvidenceItem).where(EvidenceItem.project_id == project_id)
    )
    return {evidence_item.id: evidence_item for evidence_item in evidence_items}


def get_extracted_content_by_id(
    session: Session,
    project_id: str,
) -> dict[str, ExtractedContent]:
    extracted_contents = session.scalars(
        select(ExtractedContent).where(ExtractedContent.project_id == project_id)
    ).all()
    return {
        extracted_content.id: extracted_content
        for extracted_content in extracted_contents
    }


def get_ppt_extracted_contents(
    extracted_content_by_id: dict[str, ExtractedContent],
) -> list[ExtractedContent]:
    return [
        extracted_content
        for extracted_content in extracted_content_by_id.values()
        if get_extracted_source_role(extracted_content) == "kt_ppt"
    ]


def has_plan_source_content(
    extracted_content_by_id: dict[str, ExtractedContent],
) -> bool:
    for extracted_content in extracted_content_by_id.values():
        source_role = get_extracted_source_role(extracted_content)
        if source_role in {"fdd", "kt_ppt", "kt_transcript", "kt_session"}:
            return True

    return False


def has_extracted_source_role(
    extracted_content_by_id: dict[str, ExtractedContent],
    source_role: str,
) -> bool:
    return any(
        get_extracted_source_role(extracted_content) == source_role
        for extracted_content in extracted_content_by_id.values()
    )


def resolve_plan_payload(
    session: Session,
    project_id: str,
    extracted_content_by_id: dict[str, ExtractedContent],
) -> dict[str, Any]:
    aud_plan = latest_aud_plan(session, project_id)
    plan_payload = parse_plan_json(aud_plan)
    generation_basis = plan_payload.get("generation_basis")
    has_fdd_content = has_extracted_source_role(extracted_content_by_id, "fdd")
    has_ppt_content = has_extracted_source_role(extracted_content_by_id, "kt_ppt")
    needs_fdd_refresh = (
        has_fdd_content
        and isinstance(generation_basis, str)
        and generation_basis not in {"fdd_headings", "fdd_headings_with_ppt_support"}
    )
    needs_ppt_support_refresh = (
        has_fdd_content
        and has_ppt_content
        and generation_basis == "fdd_headings"
    )
    needs_enterprise_refresh = (
        isinstance(generation_basis, str)
        and generation_basis
        in {
            "fdd_headings",
            "fdd_headings_with_ppt_support",
            "ppt_slide_titles",
            "transcript_generic_sections",
            "standard_sections_only",
        }
        and not plan_has_enterprise_structure_section(plan_payload)
    )

    if (
        plan_has_content_sections(plan_payload)
        and not needs_fdd_refresh
        and not needs_ppt_support_refresh
        and not needs_enterprise_refresh
    ):
        return plan_payload

    if not has_plan_source_content(extracted_content_by_id):
        return plan_payload

    from app.services.aud_plan_service import generate_aud_plan

    refreshed_plan = generate_aud_plan(session, project_id)
    return parse_plan_json(refreshed_plan)


def get_section_title(section: dict[str, Any]) -> str | None:
    title = section.get("title")
    return title if isinstance(title, str) else None


def get_section_id(section: dict[str, Any]) -> str | None:
    section_id = section.get("section_id")
    return section_id if isinstance(section_id, str) else None


def parse_draft_json(draft: AUDSectionDraft | None) -> dict[str, Any]:
    if draft is None:
        return {}

    return parse_json_object(draft.draft_json)


def ensure_json_list(value: Any) -> list[Any]:
    if value is None:
        return []

    return value if isinstance(value, list) else [value]


def build_section_draft_lookup(
    drafts: list[AUDSectionDraft],
) -> dict[str, AUDSectionDraft]:
    lookup: dict[str, AUDSectionDraft] = {}

    for draft in drafts:
        lookup[f"id:{normalize_title(draft.section_id)}"] = draft
        lookup[f"title:{normalize_title(draft.title)}"] = draft

    return lookup


def get_matching_section_draft(
    section: dict[str, Any],
    draft_lookup: dict[str, AUDSectionDraft],
) -> AUDSectionDraft | None:
    section_id = get_section_id(section)
    if section_id:
        draft = draft_lookup.get(f"id:{normalize_title(section_id)}")
        if draft is not None:
            return draft

    title = get_section_title(section)
    if not title:
        return None

    return draft_lookup.get(f"title:{normalize_title(title)}")


def is_draft_allowed(
    draft: AUDSectionDraft | None,
    options: DocxGenerationOptions,
) -> bool:
    if draft is None or not options.use_ai_drafts:
        return False

    review_status = normalize_title(draft.review_status)

    if review_status in ACCEPTED_DRAFT_STATUSES:
        return True

    return review_status == "draft" and options.include_draft_sections


def is_draft_omitted(draft: AUDSectionDraft | None) -> bool:
    return draft is not None and normalize_title(draft.review_status) in OMITTED_DRAFT_STATUSES


def get_section_source_role_basis(section: dict[str, Any]) -> str:
    source_role_basis = section.get("source_role_basis")
    return source_role_basis if isinstance(source_role_basis, str) else "unknown"


def get_source_content_ids(section: dict[str, Any]) -> list[str]:
    source_content_ids = section.get("source_content_ids")

    if not isinstance(source_content_ids, list):
        return []

    return [
        source_content_id
        for source_content_id in source_content_ids
        if isinstance(source_content_id, str)
    ]


def get_section_source_slide_numbers(section: dict[str, Any]) -> set[int]:
    notes = section.get("notes")

    if not isinstance(notes, list):
        return set()

    slide_numbers: set[int] = set()
    for note in notes:
        if not isinstance(note, str):
            continue

        match = re.search(r"\bSource slide\s+(\d+)\b", note, flags=re.IGNORECASE)
        if match:
            slide_numbers.add(int(match.group(1)))

    return slide_numbers


def get_extracted_source_role(extracted_content: ExtractedContent) -> str:
    json_content = parse_extracted_json(extracted_content)
    source_role = json_content.get("source_role")
    if isinstance(source_role, str) and source_role != "unknown":
        return source_role

    if extracted_content.content_type == "pptx":
        return "kt_ppt"

    return "unknown"


def get_mapped_extracted_contents(
    section: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
) -> list[ExtractedContent]:
    mapped_contents: list[ExtractedContent] = []

    for source_content_id in get_source_content_ids(section):
        extracted_content = extracted_content_by_id.get(source_content_id)
        if extracted_content is not None:
            mapped_contents.append(extracted_content)

    return mapped_contents


def meaningful_paragraphs(value: str | None) -> list[str]:
    if not value:
        return []

    paragraphs = []
    for paragraph in re.split(r"\n\s*\n", value):
        normalized_paragraph = " ".join(paragraph.split())
        if normalized_paragraph:
            paragraphs.append(normalized_paragraph)

    return paragraphs


def table_aware_paragraphs(value: str | None) -> list[str]:
    if not value:
        return []

    paragraphs: list[str] = []
    for paragraph in re.split(r"\n\s*\n", value):
        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
        if not lines:
            continue

        preserved_paragraph = "\n".join(lines)
        if looks_like_table_text(preserved_paragraph):
            paragraphs.append(preserved_paragraph)
        else:
            paragraphs.append(" ".join(preserved_paragraph.split()))

    return paragraphs


def truncate_paragraphs(
    paragraphs: list[str],
    max_paragraphs: int = MAX_SECTION_PARAGRAPHS,
    max_characters: int = MAX_SECTION_CHARACTERS,
) -> tuple[list[str], bool]:
    return paragraphs, False


def is_heading_block(block: str) -> bool:
    return bool(re.match(r"^\[Heading:\s*.+\]$", block.strip()))


def get_heading_title(block: str) -> str | None:
    match = re.match(r"^\[Heading:\s*(.+)\]$", block.strip())
    return match.group(1).strip() if match else None


def extract_fdd_heading_content(
    extracted_content: ExtractedContent,
    heading_title: str,
    max_paragraphs: int = MAX_SECTION_PARAGRAPHS,
    max_characters: int = MAX_SECTION_CHARACTERS,
) -> tuple[list[str], bool]:
    blocks = table_aware_paragraphs(extracted_content.text_content)
    collecting = False
    paragraphs: list[str] = []

    for block in blocks:
        if is_heading_block(block):
            current_heading = get_heading_title(block)
            if collecting:
                break

            collecting = normalize_title(current_heading or "") == normalize_title(
                heading_title
            )
            continue

        if collecting:
            paragraphs.append(block)

    return truncate_paragraphs(paragraphs, max_paragraphs, max_characters)


def extract_enterprise_structure_content(
    extracted_content: ExtractedContent,
) -> tuple[list[str], bool]:
    blocks = table_aware_paragraphs(extracted_content.text_content)
    collecting = False
    paragraphs: list[str] = []

    for block in blocks:
        if is_heading_block(block):
            current_heading = get_heading_title(block)
            if collecting:
                break

            collecting = is_enterprise_structure_title(current_heading)
            continue

        if not collecting and is_enterprise_structure_title(block):
            collecting = True
            continue

        if collecting:
            paragraphs.append(block)

    return truncate_paragraphs(
        paragraphs,
        MAX_ENTERPRISE_SECTION_PARAGRAPHS,
        MAX_ENTERPRISE_SECTION_CHARACTERS,
    )


def render_slide_paragraphs(slide: dict[str, Any]) -> list[str]:
    paragraphs: list[str] = []

    texts = slide.get("texts")
    if isinstance(texts, list):
        paragraphs.extend(
            text for text in texts if isinstance(text, str) and text.strip()
        )

    notes = slide.get("notes")
    if isinstance(notes, str) and notes.strip():
        paragraphs.append(notes.strip())

    return meaningful_paragraphs("\n\n".join(paragraphs))


def get_slide_image_paths(
    slide: dict[str, Any],
    fallback_image_paths: list[str],
    fallback_start_index: int,
) -> list[str]:
    image_paths = slide.get("image_paths")

    if isinstance(image_paths, list):
        return [path for path in image_paths if isinstance(path, str)]

    image_count = slide.get("image_count")
    if not isinstance(image_count, int) or image_count <= 0:
        return []

    return fallback_image_paths[fallback_start_index : fallback_start_index + image_count]


def iter_ppt_slides_with_images(
    extracted_content: ExtractedContent,
) -> list[dict[str, Any]]:
    json_content = parse_extracted_json(extracted_content)
    slides = json_content.get("slides")

    if not isinstance(slides, list):
        return []

    fallback_image_paths = json_content.get("image_paths")
    fallback_paths = (
        [path for path in fallback_image_paths if isinstance(path, str)]
        if isinstance(fallback_image_paths, list)
        else []
    )
    fallback_index = 0
    slides_with_images: list[dict[str, Any]] = []

    for slide in slides:
        if not isinstance(slide, dict):
            continue

        image_paths = get_slide_image_paths(slide, fallback_paths, fallback_index)
        image_count = slide.get("image_count")
        if isinstance(image_count, int) and image_count > 0:
            fallback_index += image_count

        if not image_paths:
            continue

        slide_copy = dict(slide)
        slide_copy["image_paths"] = image_paths
        slides_with_images.append(slide_copy)

    return slides_with_images


def is_low_value_slide(slide: dict[str, Any]) -> bool:
    title = slide.get("title")
    normalized_title = normalize_title(title if isinstance(title, str) else "")

    if normalized_title in LOW_VALUE_SLIDE_TITLES:
        return True

    if normalized_title.startswith("agenda:"):
        return True

    if "kt session" in normalized_title or "knowledge transfer" in normalized_title:
        return True

    return not normalized_title and not render_slide_paragraphs(slide)


def has_meaningful_slide_text(slide: dict[str, Any]) -> bool:
    return bool(render_slide_paragraphs(slide))


def extract_ppt_slide_content(
    extracted_content: ExtractedContent,
    slide_title: str,
    max_paragraphs: int = MAX_SECTION_PARAGRAPHS,
    max_characters: int = MAX_SECTION_CHARACTERS,
) -> tuple[list[str], bool]:
    json_content = parse_extracted_json(extracted_content)
    slides = json_content.get("slides")

    if not isinstance(slides, list):
        return [], False

    paragraphs: list[str] = []
    for slide in slides:
        if not isinstance(slide, dict):
            continue

        title = slide.get("title")
        if not isinstance(title, str):
            continue

        if normalize_title(title) == normalize_title(slide_title):
            paragraphs.extend(render_slide_paragraphs(slide))

    return truncate_paragraphs(paragraphs, max_paragraphs, max_characters)


def get_docx_extracted_contents(
    extracted_content_by_id: dict[str, ExtractedContent],
) -> list[ExtractedContent]:
    return [
        extracted_content
        for extracted_content in extracted_content_by_id.values()
        if extracted_content.content_type == "docx"
    ]


def collect_section_ppt_images(
    section: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
) -> list[dict[str, Any]]:
    section_title = get_section_title(section)
    if not section_title:
        return []

    source_slide_numbers = get_section_source_slide_numbers(section)
    mapped_ppt_content_ids = {
        content.id
        for content in get_mapped_extracted_contents(section, extracted_content_by_id)
        if get_extracted_source_role(content) == "kt_ppt"
    }
    candidates: list[dict[str, Any]] = []
    seen_paths: set[str] = set()

    for extracted_content in get_ppt_extracted_contents(extracted_content_by_id):
        for slide in iter_ppt_slides_with_images(extracted_content):
            if is_low_value_slide(slide):
                continue

            slide_title = slide.get("title")
            title_matches = isinstance(slide_title, str) and titles_resemble(
                section_title,
                slide_title,
            )
            mapped_meaningful_slide = (
                extracted_content.id in mapped_ppt_content_ids
                and slide.get("slide_number") in source_slide_numbers
                and has_meaningful_slide_text(slide)
            )

            if not title_matches and not mapped_meaningful_slide:
                continue

            slide_number = slide.get("slide_number")
            caption_title = slide_title if isinstance(slide_title, str) else "Untitled"

            for image_path in slide["image_paths"]:
                if not is_supported_docx_image_path(image_path):
                    continue

                if image_path in seen_paths:
                    continue

                seen_paths.add(image_path)
                candidates.append(
                    {
                        "storage_path": image_path,
                        "slide_number": slide_number,
                        "slide_title": caption_title,
                    }
                )

                if len(candidates) >= MAX_IMAGES_PER_SECTION:
                    return candidates

    return candidates


def collect_section_docx_images(
    section: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
) -> list[dict[str, Any]]:
    section_title = get_section_title(section)
    if not section_title:
        return []

    mapped_docx_content_ids = {
        content.id
        for content in get_mapped_extracted_contents(section, extracted_content_by_id)
        if content.content_type == "docx"
    }
    include_all_docx_sources = is_enterprise_structure_title(section_title)
    candidates: list[dict[str, Any]] = []
    seen_paths: set[str] = set()

    for extracted_content in get_docx_extracted_contents(extracted_content_by_id):
        if (
            not include_all_docx_sources
            and extracted_content.id not in mapped_docx_content_ids
        ):
            continue

        json_content = parse_extracted_json(extracted_content)
        images = json_content.get("images")
        if not isinstance(images, list):
            continue

        for image in images:
            if not isinstance(image, dict):
                continue

            image_section_title = image.get("section_title")
            if not (
                isinstance(image_section_title, str)
                and titles_resemble(section_title, image_section_title)
            ):
                continue

            storage_path = image.get("storage_path")
            if not isinstance(storage_path, str):
                continue

            if not is_supported_docx_image_path(storage_path):
                continue

            if storage_path in seen_paths:
                continue

            seen_paths.add(storage_path)
            candidates.append(
                {
                    "storage_path": storage_path,
                    "slide_number": None,
                    "slide_title": image_section_title,
                    "caption": f"Source image from DOCX section: {image_section_title}",
                }
            )

    return candidates


def collect_section_images(
    section: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
) -> list[dict[str, Any]]:
    images = [
        *collect_section_docx_images(section, extracted_content_by_id),
        *collect_section_ppt_images(section, extracted_content_by_id),
    ]

    if is_enterprise_structure_title(get_section_title(section)):
        return images

    return images[:MAX_IMAGES_PER_SECTION]


def is_supported_docx_image_path(storage_path: str) -> bool:
    return Path(storage_path).suffix.lower() in SUPPORTED_DOCX_IMAGE_EXTENSIONS


def resolve_storage_service(
    storage_root: Path | None = None,
    storage_service: StorageService | None = None,
) -> StorageService:
    if storage_service is not None:
        return storage_service

    if storage_root is not None:
        return LocalStorageService(storage_root)

    return get_file_storage()


def materialize_image_path(
    storage_service: StorageService,
    storage_path: str,
    temporary_dir: Path,
) -> Path | None:
    local_path = storage_service.local_path(storage_path)
    if local_path is not None:
        return local_path

    destination = temporary_dir / Path(storage_path).name
    try:
        storage_service.download_to_path(storage_path, destination)
    except Exception:
        return None

    return destination if destination.is_file() else None


def get_usable_page_width(document: Document) -> Emu:
    section = document.sections[-1]
    return Emu(section.page_width - section.left_margin - section.right_margin)


def add_section_images(
    document: Document,
    images: list[dict[str, Any]],
    storage_service: StorageService,
    temporary_dir: Path,
) -> None:
    if not images:
        return

    usable_width = get_usable_page_width(document)

    for image in images:
        storage_path = image.get("storage_path")
        if not isinstance(storage_path, str):
            continue

        image_path = materialize_image_path(
            storage_service,
            storage_path,
            temporary_dir,
        )
        if image_path is None:
            continue

        try:
            document.add_picture(str(image_path), width=usable_width)
        except (OSError, UnrecognizedImageError):
            continue

        slide_number = image.get("slide_number")
        slide_title = image.get("slide_title")
        custom_caption = image.get("caption")
        if isinstance(custom_caption, str):
            caption = custom_caption
        else:
            caption = (
                f"Source image from slide {slide_number}: {slide_title}"
                if isinstance(slide_number, int)
                else f"Source image from slide: {slide_title}"
            )
        caption_paragraph = document.add_paragraph(clean_generated_text(caption))
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_evidence_json(evidence_item: EvidenceItem | None) -> dict[str, Any]:
    if evidence_item is None:
        return {}

    return parse_json_object(evidence_item.json_data)


def normalize_table_rows(rows: Any) -> list[list[str]]:
    if not isinstance(rows, list):
        return []

    normalized_rows: list[list[str]] = []
    for row in rows:
        if isinstance(row, list):
            normalized_rows.append([str(cell or "") for cell in row])
        elif isinstance(row, dict):
            values = row.get("values")
            if isinstance(values, list):
                normalized_rows.append([str(cell or "") for cell in values])

    return [row for row in normalized_rows if any(cell.strip() for cell in row)]


def rows_from_table_item(
    table_item: Any,
    evidence_item_by_id: dict[str, EvidenceItem],
) -> tuple[str | None, list[list[str]]]:
    if isinstance(table_item, str):
        evidence_item = evidence_item_by_id.get(table_item)
        return evidence_item.title if evidence_item else None, rows_from_evidence(evidence_item)

    if not isinstance(table_item, dict):
        return None, []

    title = table_item.get("title")
    table_title = title if isinstance(title, str) else None
    rows = normalize_table_rows(table_item.get("rows"))

    if rows:
        return table_title, rows

    evidence_item_id = table_item.get("evidence_item_id") or table_item.get("id")
    if not isinstance(evidence_item_id, str):
        return table_title, []

    evidence_item = evidence_item_by_id.get(evidence_item_id)
    return table_title or (evidence_item.title if evidence_item else None), rows_from_evidence(
        evidence_item
    )


def rows_from_evidence(evidence_item: EvidenceItem | None) -> list[list[str]]:
    if evidence_item is None:
        return []

    json_data = parse_evidence_json(evidence_item)

    for key_path in (
        ("table", "rows"),
        ("slide", "tables"),
        ("rows",),
        ("sheet", "rows"),
    ):
        current: Any = json_data
        for key in key_path:
            if isinstance(current, dict):
                current = current.get(key)
            else:
                current = None
                break

        if key_path == ("slide", "tables") and isinstance(current, list):
            for table in current:
                if isinstance(table, dict):
                    rows = normalize_table_rows(table.get("rows"))
                    if rows:
                        return rows

        rows = normalize_table_rows(current)
        if rows:
            return rows

    text_rows = [
        [cell.strip() for cell in line.split("|")]
        for line in (evidence_item.text or "").splitlines()
        if line.strip()
    ]
    return [row for row in text_rows if any(cell for cell in row)]


def add_section_tables(
    document: Document,
    included_tables: list[Any],
    evidence_item_by_id: dict[str, EvidenceItem],
) -> None:
    for table_item in included_tables:
        title, rows = rows_from_table_item(table_item, evidence_item_by_id)
        if not rows:
            continue

        if title:
            document.add_paragraph(clean_generated_text(title))

        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=1, cols=column_count)
        table.style = "Table Grid"

        for index, value in enumerate(rows[0]):
            table.rows[0].cells[index].text = clean_generated_text(value)

        for source_row in rows[1:]:
            row = table.add_row().cells
            for index, value in enumerate(source_row[:column_count]):
                row[index].text = clean_generated_text(value)


def image_from_evidence(
    evidence_item: EvidenceItem | None,
) -> dict[str, Any] | None:
    if evidence_item is None:
        return None

    json_data = parse_evidence_json(evidence_item)
    image_path = json_data.get("image_path")
    slide_number = json_data.get("slide_number")

    if not isinstance(image_path, str):
        image_path = evidence_item.text if isinstance(evidence_item.text, str) else None

    if not isinstance(image_path, str):
        return None

    return {
        "storage_path": image_path,
        "slide_number": slide_number if isinstance(slide_number, int) else None,
        "slide_title": evidence_item.title or "Selected image",
    }


def normalize_included_images(
    included_images: list[Any],
    evidence_item_by_id: dict[str, EvidenceItem],
) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    seen_paths: set[str] = set()

    for item in included_images:
        image: dict[str, Any] | None = None

        if isinstance(item, str):
            evidence_image = image_from_evidence(evidence_item_by_id.get(item))
            image = evidence_image or {"storage_path": item, "slide_title": "Selected image"}
        elif isinstance(item, dict):
            storage_path = (
                item.get("storage_path")
                or item.get("image_path")
                or item.get("path")
            )
            if isinstance(storage_path, str):
                image = {
                    "storage_path": storage_path,
                    "slide_number": item.get("slide_number"),
                    "slide_title": item.get("slide_title") or item.get("title") or "Selected image",
                    "caption": item.get("caption"),
                }
            else:
                evidence_item_id = item.get("evidence_item_id") or item.get("id")
                if isinstance(evidence_item_id, str):
                    image = image_from_evidence(evidence_item_by_id.get(evidence_item_id))

        if image is None:
            continue

        storage_path = image.get("storage_path")
        if not isinstance(storage_path, str):
            continue

        if not is_supported_docx_image_path(storage_path) or storage_path in seen_paths:
            continue

        seen_paths.add(storage_path)
        images.append(image)

    return images


def style_hint_for_section(section_title: str | None) -> str:
    normalized = normalize_title(section_title or "")

    if "open point" in normalized:
        return "open_points"
    if "activity" in normalized:
        return "activity"
    if "ricew" in normalized:
        return "ricew"
    if "report" in normalized:
        return "reports"
    if "feature" in normalized and "enable" in normalized:
        return "feature_enablement"

    return "standard"


def to_populated_table(normalized_table: NormalizedTable) -> PopulatedTable:
    return PopulatedTable(
        title=clean_generated_text(normalized_table.title),
        columns=[clean_generated_text(column) for column in normalized_table.columns],
        rows=[
            [clean_generated_text(cell) for cell in row]
            for row in normalized_table.rows
        ],
        source=normalized_table.source,
        section_id=normalized_table.section_id,
        style_hint=normalized_table.style_hint,
    )


def looks_like_table_text(value: str) -> bool:
    stripped = value.strip()
    if "|" not in stripped:
        return False

    if re.match(r"^\[?Table\s+\d+[:\]]?", stripped, flags=re.IGNORECASE):
        return True

    return len([line for line in stripped.splitlines() if "|" in line]) >= 2


def extract_structured_tables_from_paragraphs(
    paragraphs: list[str],
    *,
    section_title: str | None,
    section_id: str | None,
    source: str | None,
    drop_unparsed_table_text: bool = False,
) -> tuple[list[str], list[PopulatedTable]]:
    remaining_paragraphs: list[str] = []
    tables: list[PopulatedTable] = []
    style_hint = style_hint_for_section(section_title)

    for paragraph in paragraphs:
        if not looks_like_table_text(paragraph):
            remaining_paragraphs.append(paragraph)
            continue

        if drop_unparsed_table_text:
            continue

        result = TABLE_NORMALIZER.normalize_with_reason(
            paragraph,
            title=None,
            source=source,
            section_id=section_id,
            style_hint=style_hint,
        )
        if result.table is None:
            logger.warning(
                "DOCX table fallback for section %s: %s",
                section_title or "unknown",
                result.fallback_reason,
            )
            remaining_paragraphs.append(paragraph)
            continue

        tables.append(to_populated_table(result.table))

    return remaining_paragraphs, tables


def collect_docx_section_tables(
    extracted_content: ExtractedContent,
    section_title: str,
    section_id: str | None,
) -> list[PopulatedTable]:
    json_content = parse_extracted_json(extracted_content)
    tables = json_content.get("tables")
    if not isinstance(tables, list):
        return []

    populated_tables: list[PopulatedTable] = []
    for table in tables:
        if not isinstance(table, dict):
            continue

        table_section_title = table.get("section_title")
        if isinstance(table_section_title, str) and not titles_resemble(
            section_title,
            table_section_title,
        ):
            continue

        table_index = table.get("index")
        table_title = (
            f"Table {table_index}: {section_title}"
            if table_index
            else section_title
        )
        populated_table = normalize_source_table(
            table,
            title=table_title,
            source=extracted_content.title,
            section_id=section_id,
            style_hint=style_hint_for_section(section_title),
        )
        if populated_table is not None:
            populated_tables.append(populated_table)

    return populated_tables


def normalize_source_table(
    table_data: Any,
    *,
    title: str | None,
    source: str | None,
    section_id: str | None,
    style_hint: str,
) -> PopulatedTable | None:
    result = TABLE_NORMALIZER.normalize_with_reason(
        table_data,
        title=title,
        source=source,
        section_id=section_id,
        style_hint=style_hint,
    )
    if result.table is None:
        logger.warning(
            "DOCX table fallback for %s: %s",
            title or source or "source table",
            result.fallback_reason,
        )
        return None

    return to_populated_table(result.table)


def extract_ppt_slide_tables(
    extracted_content: ExtractedContent,
    slide_title: str,
    section_id: str | None,
) -> list[PopulatedTable]:
    json_content = parse_extracted_json(extracted_content)
    slides = json_content.get("slides")

    if not isinstance(slides, list):
        return []

    populated_tables: list[PopulatedTable] = []
    for slide in slides:
        if not isinstance(slide, dict):
            continue

        title = slide.get("title")
        if not isinstance(title, str) or normalize_title(title) != normalize_title(slide_title):
            continue

        tables = slide.get("tables")
        if not isinstance(tables, list):
            continue

        for table in tables:
            if not isinstance(table, dict):
                continue

            table_index = table.get("index")
            table_title = (
                f"Slide {slide.get('slide_number')} Table {table_index}"
                if table_index
                else f"Slide {slide.get('slide_number')} Table"
            )
            populated_table = normalize_source_table(
                table,
                title=table_title,
                source=extracted_content.title,
                section_id=section_id,
                style_hint=style_hint_for_section(slide_title),
            )
            if populated_table is not None:
                populated_tables.append(populated_table)

    return populated_tables


def build_section_content(
    section: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
) -> tuple[list[str], list[PopulatedTable], bool]:
    section_title = get_section_title(section)
    section_id = get_section_id(section)

    if not section_title:
        return [], [], False

    mapped_contents = get_mapped_extracted_contents(section, extracted_content_by_id)
    fdd_contents = [
        content
        for content in mapped_contents
        if get_extracted_source_role(content) == "fdd"
    ]
    source_role_basis = get_section_source_role_basis(section)
    is_enterprise_section = is_enterprise_structure_title(section_title)

    if is_enterprise_section:
        enterprise_contents = mapped_contents or list(extracted_content_by_id.values())
        for extracted_content in enterprise_contents:
            if extracted_content.content_type == "docx":
                paragraphs, was_truncated = extract_enterprise_structure_content(
                    extracted_content
                )
                if paragraphs:
                    docx_tables = collect_docx_section_tables(
                        extracted_content,
                        section_title,
                        section_id,
                    )
                    paragraphs, parsed_tables = extract_structured_tables_from_paragraphs(
                        paragraphs,
                        section_title=section_title,
                        section_id=section_id,
                        source=extracted_content.title,
                        drop_unparsed_table_text=bool(docx_tables),
                    )
                    return paragraphs, [*docx_tables, *parsed_tables], was_truncated
            if get_extracted_source_role(extracted_content) == "kt_ppt":
                paragraphs, was_truncated = extract_ppt_slide_content(
                    extracted_content,
                    ENTERPRISE_STRUCTURE_TITLE,
                    MAX_ENTERPRISE_SECTION_PARAGRAPHS,
                    MAX_ENTERPRISE_SECTION_CHARACTERS,
                )
                if paragraphs:
                    return (
                        paragraphs,
                        extract_ppt_slide_tables(
                            extracted_content,
                            ENTERPRISE_STRUCTURE_TITLE,
                            section_id,
                        ),
                        was_truncated,
                    )

    for extracted_content in fdd_contents:
        paragraphs, was_truncated = extract_fdd_heading_content(
            extracted_content,
            section_title,
            MAX_ENTERPRISE_SECTION_PARAGRAPHS
            if is_enterprise_section
            else MAX_SECTION_PARAGRAPHS,
            MAX_ENTERPRISE_SECTION_CHARACTERS
            if is_enterprise_section
            else MAX_SECTION_CHARACTERS,
        )
        if paragraphs:
            docx_tables = (
                collect_docx_section_tables(extracted_content, section_title, section_id)
                if extracted_content.content_type == "docx"
                else []
            )
            paragraphs, parsed_tables = extract_structured_tables_from_paragraphs(
                paragraphs,
                section_title=section_title,
                section_id=section_id,
                source=extracted_content.title,
                drop_unparsed_table_text=bool(docx_tables),
            )
            return paragraphs, [*docx_tables, *parsed_tables], was_truncated

    if source_role_basis == "fdd":
        return [], [], False

    if source_role_basis == "kt_ppt":
        for extracted_content in mapped_contents:
            if get_extracted_source_role(extracted_content) != "kt_ppt":
                continue

            paragraphs, was_truncated = extract_ppt_slide_content(
                extracted_content,
                section_title,
                MAX_ENTERPRISE_SECTION_PARAGRAPHS
                if is_enterprise_section
                else MAX_SECTION_PARAGRAPHS,
                MAX_ENTERPRISE_SECTION_CHARACTERS
                if is_enterprise_section
                else MAX_SECTION_CHARACTERS,
            )
            if paragraphs:
                return (
                    paragraphs,
                    extract_ppt_slide_tables(
                        extracted_content,
                        section_title,
                        section_id,
                    ),
                    was_truncated,
                )

    return [], [], False


def build_section_render_context(
    section: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
    draft_lookup: dict[str, AUDSectionDraft],
    options: DocxGenerationOptions,
) -> SectionRenderContext | None:
    draft = get_matching_section_draft(section, draft_lookup)

    if is_draft_omitted(draft):
        return None

    draft_json = parse_draft_json(draft)
    if is_draft_allowed(draft, options):
        paragraphs = table_aware_paragraphs(draft.draft_text)
        if paragraphs and paragraphs != [PLACEHOLDER_LINE]:
            paragraphs, tables = extract_structured_tables_from_paragraphs(
                paragraphs,
                section_title=get_section_title(section),
                section_id=get_section_id(section),
                source=draft.title,
            )
            return SectionRenderContext(
                draft=draft,
                draft_json=draft_json,
                paragraphs=paragraphs,
                tables=tables,
                was_truncated=False,
            )

    paragraphs, tables, was_truncated = build_section_content(
        section,
        extracted_content_by_id,
    )
    return SectionRenderContext(
        draft=draft,
        draft_json=draft_json,
        paragraphs=paragraphs,
        tables=tables,
        was_truncated=was_truncated,
    )


def get_final_plan_sections(plan_payload: dict[str, Any]) -> list[Any]:
    ai_enhanced_plan = plan_payload.get("ai_enhanced_plan")
    if isinstance(ai_enhanced_plan, dict) and isinstance(
        ai_enhanced_plan.get("sections"),
        list,
    ):
        return ai_enhanced_plan["sections"]

    sections = plan_payload.get("sections")
    return sections if isinstance(sections, list) else []


def iter_planned_sections(plan_payload: dict[str, Any]) -> list[dict[str, Any]]:
    sections = get_final_plan_sections(plan_payload)

    if not isinstance(sections, list):
        return []

    planned_sections: list[dict[str, Any]] = []
    seen_titles: set[str] = set()

    for section in sections:
        if not isinstance(section, dict):
            continue

        title = section.get("title")
        include_in_aud = section.get("include_in_aud", True)
        missing_info_handling = section.get("missing_info_handling")

        if not isinstance(title, str) or include_in_aud is False:
            continue

        normalized_title = normalize_title(title)
        if (
            isinstance(missing_info_handling, str)
            and normalize_title(missing_info_handling) == "omit"
        ):
            continue

        if (
            normalized_title in SKIP_PLANNED_SECTION_TITLES
            or normalized_title in seen_titles
        ):
            continue

        seen_titles.add(normalized_title)
        planned_sections.append(section)

    if not planned_sections:
        return planned_sections

    if any(
        is_enterprise_structure_title(get_section_title(section))
        for section in planned_sections
    ):
        return planned_sections

    enterprise_section = {
        "title": ENTERPRISE_STRUCTURE_TITLE,
        "include_in_aud": True,
        "source_role_basis": "required_placeholder",
        "source_content_ids": [],
        "notes": ["Required carry-forward section."],
    }
    insert_index = 0
    for index, section in enumerate(planned_sections):
        if normalize_title(section["title"]) == "introduction":
            insert_index = index + 1
            break

    return [
        *planned_sections[:insert_index],
        enterprise_section,
        *planned_sections[insert_index:],
    ]


def format_table_of_contents(titles: list[str]) -> str:
    return "\n".join(
        f"{index}. {title}" for index, title in enumerate(titles, start=1)
    )


def build_template_replacements(
    document_model: PopulatedDocumentModel,
) -> dict[str, str]:
    module_name = document_model.module_name
    customer_name = document_model.customer_name
    author_name = document_model.author_name
    table_of_contents = format_table_of_contents(document_model.toc_titles)
    open_points_text = "\n".join(
        f"{index}. {open_point.question}"
        for index, open_point in enumerate(document_model.open_points, start=1)
    )

    return {
        "<Customer Name>": customer_name,
        "<Module Name>": module_name,
        "<Date>": document_model.generated_date.isoformat(),
        "<Author>": author_name,
        "<Ver.>": document_model.version_history.version,
        "<Table Of Content>": table_of_contents,
        "<Table of Content>": table_of_contents,
        "<Table of Contents>": table_of_contents,
        "<From User Input>": module_name,
        "<Generate from Customer name>": customer_name,
        "<Process Name>": module_name,
        "<Subprocess Name>": "",
        "<Attributes>": "",
        "<Open Points>": open_points_text,
        "<Unsresolved or missing information to be listed here>": open_points_text,
    }


def apply_replacements_to_text(text: str, replacements: dict[str, str]) -> str:
    replaced_text = text
    for placeholder, value in replacements.items():
        replaced_text = replaced_text.replace(placeholder, value)

    return TEMPLATE_PLACEHOLDER_PATTERN.sub("", replaced_text)


def replace_paragraph_placeholders(
    paragraph,
    replacements: dict[str, str],
) -> None:
    original_text = paragraph.text
    replaced_text = apply_replacements_to_text(original_text, replacements)

    if replaced_text == original_text:
        return

    if not paragraph.runs:
        paragraph.add_run(replaced_text)
        return

    paragraph.runs[0].text = replaced_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_template_placeholders(
    document: Document,
    document_model: PopulatedDocumentModel,
) -> None:
    replacements = build_template_replacements(document_model)

    for paragraph in document.paragraphs:
        replace_paragraph_placeholders(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_placeholders(paragraph, replacements)


def apply_generated_heading_format(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(SECTION_HEADING_SPACE_BEFORE_PT)
    paragraph_format.space_after = Pt(SECTION_HEADING_SPACE_AFTER_PT)
    paragraph_format.keep_with_next = True

    for run in paragraph.runs:
        run.font.name = SECTION_HEADING_FONT_NAME
        run.font.size = Pt(SECTION_HEADING_FONT_SIZE_PT)
        run.font.bold = True
        run.font.color.rgb = SECTION_HEADING_COLOR
        run._element.rPr.rFonts.set(qn("w:eastAsia"), SECTION_HEADING_FONT_NAME)


def add_generated_heading(
    document: Document,
    title: str,
    level: int = 1,
):
    heading = document.add_heading(title, level=level)
    apply_generated_heading_format(heading)
    return heading


def truncate_template_after_cover(document: Document) -> None:
    body = document._body._element
    body_children = list(body)
    start_element = None

    for paragraph in document.paragraphs:
        if normalize_title(paragraph.text) in TEMPLATE_BODY_START_TITLES:
            start_element = paragraph._element
            break

    if start_element is None:
        return

    try:
        start_index = body_children.index(start_element)
    except ValueError:
        return

    for child in body_children[start_index:]:
        if child.tag.endswith("}sectPr"):
            continue

        body.remove(child)


def add_table_of_contents(
    document: Document,
    document_model: PopulatedDocumentModel,
) -> None:
    add_generated_heading(document, "Table of Contents")
    TABLE_RENDERER.add_table(
        document,
        NormalizedTable(
            title=None,
            columns=["Section", "Title"],
            rows=[
                [str(index), title]
                for index, title in enumerate(document_model.toc_titles, start=1)
            ],
            style_hint="standard",
        ),
    )


def add_document_version_history(
    document: Document,
    document_model: PopulatedDocumentModel,
) -> None:
    add_generated_heading(document, "Document Version History")
    TABLE_RENDERER.add_table(
        document,
        NormalizedTable(
            title=None,
            columns=["Version", "Date", "Author", "Reviewed By"],
            rows=[
                [
                    document_model.version_history.version,
                    document_model.version_history.date.isoformat(),
                    document_model.version_history.author,
                    document_model.version_history.reviewed_by,
                ]
            ],
            style_hint="standard",
        ),
    )


def add_open_points_table(
    document: Document,
    open_points: list[PopulatedOpenPoint],
) -> None:
    if not open_points:
        return

    add_generated_heading(document, "Open Points")
    TABLE_RENDERER.add_table(
        document,
        NormalizedTable(
            title=None,
            columns=["ID", "Topic", "Question", "Status"],
            rows=[
                [str(index), open_point.topic, open_point.question, open_point.status]
                for index, open_point in enumerate(open_points, start=1)
            ],
            style_hint="open_points",
        ),
    )


def to_populated_images(images: list[dict[str, Any]]) -> list[PopulatedImage]:
    populated_images: list[PopulatedImage] = []

    for image in images:
        storage_path = image.get("storage_path")
        if not isinstance(storage_path, str):
            continue

        slide_number = image.get("slide_number")
        slide_title = image.get("slide_title")
        caption = image.get("caption")
        populated_images.append(
            PopulatedImage(
                storage_path=storage_path,
                slide_number=slide_number if isinstance(slide_number, int) else None,
                slide_title=clean_generated_text(slide_title)
                if isinstance(slide_title, str)
                else None,
                caption=clean_generated_text(caption)
                if isinstance(caption, str)
                else None,
            )
        )

    return populated_images


def to_image_dicts(images: list[PopulatedImage]) -> list[dict[str, Any]]:
    return [
        {
            "storage_path": image.storage_path,
            "slide_number": image.slide_number,
            "slide_title": image.slide_title,
            "caption": image.caption,
        }
        for image in images
    ]


def to_populated_tables(
    table_items: list[Any],
    evidence_item_by_id: dict[str, EvidenceItem],
    section_title: str | None = None,
    section_id: str | None = None,
) -> list[PopulatedTable]:
    tables: list[PopulatedTable] = []

    for table_item in table_items:
        title, rows = rows_from_table_item(table_item, evidence_item_by_id)
        normalized_table = TABLE_NORMALIZER.normalize(
            {
                "title": title,
                "rows": rows,
                "source": title,
                "section_id": section_id,
                "style_hint": style_hint_for_section(section_title),
            }
        )
        if normalized_table is not None:
            tables.append(to_populated_table(normalized_table))
        elif rows:
            logger.warning(
                "DOCX table fallback for selected table %s in section %s",
                title or "untitled",
                section_title or "unknown",
            )

    return tables


def build_section_population_inputs(
    plan_payload: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
    section_drafts: list[AUDSectionDraft],
    evidence_item_by_id: dict[str, EvidenceItem],
    options: DocxGenerationOptions,
) -> list[SectionPopulationInput]:
    draft_lookup = build_section_draft_lookup(section_drafts)
    section_inputs: list[SectionPopulationInput] = []

    for section in iter_planned_sections(plan_payload):
        render_context = build_section_render_context(
            section,
            extracted_content_by_id,
            draft_lookup,
            options,
        )
        if render_context is None:
            continue

        title = get_section_title(section)
        if not title:
            continue

        paragraphs = [
            clean_generated_text(paragraph)
            for paragraph in render_context.paragraphs
        ]
        paragraphs = [paragraph for paragraph in paragraphs if paragraph]
        paragraphs = dedupe_consecutive_text(paragraphs)
        tables = to_populated_tables(
            ensure_json_list(render_context.draft_json.get("included_tables")),
            evidence_item_by_id,
            section_title=title,
            section_id=get_section_id(section),
        )
        tables = [*render_context.tables, *tables]

        selected_images = normalize_included_images(
            ensure_json_list(render_context.draft_json.get("included_images")),
            evidence_item_by_id,
        )
        images = (
            to_populated_images(
                selected_images
                or (
                    collect_section_images(section, extracted_content_by_id)
                    if options.include_images
                    else []
                )
            )
            if options.include_images
            else []
        )

        section_inputs.append(
            SectionPopulationInput(
                title=title,
                paragraphs=paragraphs,
                tables=tables,
                images=images,
                source_role_basis=get_section_source_role_basis(section),
            )
        )

    return section_inputs


def add_populated_table(document: Document, populated_table: PopulatedTable) -> None:
    if not populated_table.rows:
        return

    table_data = {
        "title": populated_table.title,
        "columns": populated_table.columns,
        "rows": populated_table.rows,
        "source": populated_table.source,
        "section_id": populated_table.section_id,
        "style_hint": populated_table.style_hint,
    }
    rendered_table = TABLE_RENDERER.add_table(document, table_data)
    if rendered_table is None:
        logger.warning(
            "DOCX table fallback while rendering populated table %s",
            populated_table.title or "untitled",
        )


def set_paragraph_spacing(paragraph, *, before: int = 0, after: int = BODY_SPACE_AFTER_PT) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def is_process_flow_subheading(text: str) -> bool:
    return normalize_title(text) == "process flow"


def is_step_heading(text: str) -> bool:
    return bool(re.match(r"^step\s+\d+\s*[:.-]\s+.+", text.strip(), flags=re.IGNORECASE))


def is_step_bullet_candidate(text: str) -> bool:
    normalized = normalize_title(text)
    if not normalized:
        return False

    if is_step_heading(text) or is_process_flow_subheading(text):
        return False

    if normalized.startswith(("navigation:", "table ")):
        return False

    return True


def add_body_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    set_paragraph_spacing(paragraph)
    return paragraph


def add_content_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=6, after=4)
    run = paragraph.add_run(text)
    run.font.name = SECTION_HEADING_FONT_NAME
    run.font.size = Pt(CONTENT_SUBHEADING_FONT_SIZE_PT)
    run.font.bold = True
    run.font.color.rgb = CONTENT_SUBHEADING_COLOR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), SECTION_HEADING_FONT_NAME)


def add_step_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=8, after=3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.font.name = SECTION_HEADING_FONT_NAME
    run.font.size = Pt(STEP_HEADING_FONT_SIZE_PT)
    run.font.bold = True
    run.font.color.rgb = STEP_HEADING_COLOR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), SECTION_HEADING_FONT_NAME)


def add_readable_paragraph(document: Document, text: str) -> None:
    cleaned = clean_generated_text(text)
    if not cleaned:
        return

    bullet_match = re.match(r"^[-*]\s+(.+)$", cleaned)
    if bullet_match:
        add_bullet_paragraph(document, bullet_match.group(1))
        return

    add_body_paragraph(document, cleaned)


def add_bullet_paragraph(document: Document, text: str) -> None:
    try:
        paragraph = document.add_paragraph(text, style="List Bullet")
        set_paragraph_spacing(paragraph, after=1)
        return
    except KeyError:
        pass

    paragraph = document.add_paragraph(text)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(1)
    paragraph_format.left_indent = Pt(18)
    paragraph_format.first_line_indent = Pt(-9)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def add_section_paragraphs(document: Document, paragraphs: list[str]) -> None:
    in_step_list = False

    for paragraph in paragraphs:
        cleaned = clean_generated_text(paragraph)
        if not cleaned:
            continue

        if is_process_flow_subheading(cleaned):
            add_content_subheading(document, cleaned)
            in_step_list = False
            continue

        if is_step_heading(cleaned):
            add_step_heading(document, cleaned)
            in_step_list = True
            continue

        if in_step_list and is_step_bullet_candidate(cleaned):
            add_bullet_paragraph(document, cleaned)
            continue

        add_readable_paragraph(document, cleaned)
        in_step_list = False


def add_populated_section(
    document: Document,
    section: PopulatedSection,
    storage_service: StorageService,
    temporary_dir: Path,
) -> None:
    add_generated_heading(document, section.title)

    add_section_paragraphs(document, section.paragraphs)

    for table in section.tables:
        add_populated_table(document, table)

    add_section_images(
        document,
        to_image_dicts(section.images),
        storage_service,
        temporary_dir,
    )


def add_source_conflict_summary_appendix(
    document: Document,
    plan_payload: dict[str, Any],
) -> None:
    ai_plan = plan_payload.get("ai_enhanced_plan")
    warnings: list[str] = []

    if isinstance(ai_plan, dict):
        raw_warnings = ai_plan.get("warnings")
        if isinstance(raw_warnings, list):
            warnings.extend(str(warning) for warning in raw_warnings if str(warning))

    if not warnings:
        return

    add_generated_heading(document, "Source Conflict Summary")
    for warning in warnings:
        document.add_paragraph(warning)


def build_document(
    project: Project,
    plan_payload: dict[str, Any],
    extracted_content_by_id: dict[str, ExtractedContent],
    section_drafts: list[AUDSectionDraft],
    evidence_item_by_id: dict[str, EvidenceItem],
    open_points: list[OpenPoint],
    storage_service: StorageService,
    temporary_dir: Path,
    options: DocxGenerationOptions,
    settings: Settings,
    template_path: Path,
) -> Document:
    document = Document(str(template_path))
    current_date = utc_now().date()
    section_inputs = build_section_population_inputs(
        plan_payload=plan_payload,
        extracted_content_by_id=extracted_content_by_id,
        section_drafts=section_drafts,
        evidence_item_by_id=evidence_item_by_id,
        options=options,
    )
    document_model = TemplatePopulationService(
        selected_template_path=template_path,
        project=project,
        final_plan=plan_payload,
        section_drafts=section_drafts,
        section_inputs=section_inputs,
        open_points=open_points,
        generated_date=current_date,
    ).build_document_model()

    replace_template_placeholders(
        document=document,
        document_model=document_model,
    )
    truncate_template_after_cover(document)
    document.add_page_break()
    add_table_of_contents(document, document_model)
    add_document_version_history(document, document_model)

    add_populated_section(
        document,
        document_model.purpose_scope,
        storage_service,
        temporary_dir,
    )

    for section in document_model.sections:
        add_populated_section(
            document,
            section,
            storage_service,
            temporary_dir,
        )

    if options.include_open_points:
        add_open_points_table(document, document_model.open_points)

    if settings.INTERNAL_DEBUG_OUTPUT:
        add_source_conflict_summary_appendix(document, plan_payload)

    document.add_paragraph(GENERATION_NOTE)
    return document


def generate_docx(
    session: Session,
    project_id: str,
    storage_root: Path | None = None,
    storage_service: StorageService | None = None,
    options: DocxGenerationOptions | None = None,
    settings: Settings | None = None,
) -> GeneratedDocument:
    project = session.get(Project, project_id)

    if project is None:
        raise ValueError("Project not found.")

    resolved_options = options or DocxGenerationOptions()
    resolved_settings = settings or get_settings()
    resolved_storage_service = resolve_storage_service(storage_root, storage_service)
    extracted_content_by_id = get_extracted_content_by_id(session, project_id)
    plan_payload = resolve_plan_payload(
        session,
        project_id,
        extracted_content_by_id,
    )
    section_drafts = list_section_drafts(session, project_id)
    evidence_item_by_id = get_evidence_item_by_id(session, project_id)
    open_point_selection = list_open_points_for_docx(
        session,
        project_id,
        resolved_settings,
    )
    timestamp = utc_now().strftime("%Y%m%d-%H%M%S")
    filename_prefix = sanitize_filename_part(
        project.module_name or project.customer_name
    )
    filename = f"{filename_prefix}-aud-v1-{timestamp}.docx"
    storage_path = f"projects/{project_id}/outputs/{filename}"

    with TemporaryDirectory() as temporary_dir:
        temporary_root = Path(temporary_dir)
        template = TemplateResolver(
            session=session,
            project_id=project_id,
            storage_service=resolved_storage_service,
            settings=resolved_settings,
        ).resolve(temporary_root)
        document = build_document(
            project=project,
            plan_payload=plan_payload,
            extracted_content_by_id=extracted_content_by_id,
            section_drafts=section_drafts,
            evidence_item_by_id=evidence_item_by_id,
            open_points=open_point_selection.open_points,
            storage_service=resolved_storage_service,
            temporary_dir=temporary_root,
            options=resolved_options,
            settings=resolved_settings,
            template_path=template.path,
        )
        output_path = temporary_root / filename
        document.save(output_path)
        resolved_storage_service.write_file(storage_path, output_path)

    generated_document = GeneratedDocument(
        project_id=project_id,
        filename=filename,
        storage_path=storage_path,
        document_type=DOCUMENT_TYPE,
        metadata_json=json.dumps(
            {"open_points_fallback": open_point_selection.used_fallback}
        ),
    )
    session.add(generated_document)
    session.commit()
    session.refresh(generated_document)
    return generated_document
