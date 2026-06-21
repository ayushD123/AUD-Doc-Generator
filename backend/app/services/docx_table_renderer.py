from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

VALID_STYLE_HINTS = {
    "standard",
    "activity",
    "open_points",
    "ricew",
    "reports",
    "feature_enablement",
}
PRESET_COLUMNS = {
    "activity": ["Activity", "Owner", "Automatic/Manual", "System", "Description"],
    "open_points": ["ID", "Topic", "Question", "Status"],
    "reports": ["Report Name", "Report Description", "Module"],
    "ricew": [
        "ID",
        "Type",
        "Name",
        "Status",
        "Initiating Event",
        "Notification Type",
        "Description",
    ],
}
HEADER_FILL_BY_STYLE = {
    "standard": "D9EAF7",
    "activity": "1F4E79",
    "open_points": "806000",
    "reports": "548235",
    "ricew": "7030A0",
    "feature_enablement": "5B9BD5",
}
HEADER_TEXT_BY_STYLE = {
    "standard": RGBColor(0, 0, 0),
    "activity": RGBColor(255, 255, 255),
    "open_points": RGBColor(255, 255, 255),
    "reports": RGBColor(255, 255, 255),
    "ricew": RGBColor(255, 255, 255),
    "feature_enablement": RGBColor(255, 255, 255),
}


@dataclass(frozen=True)
class NormalizedTable:
    title: str | None
    columns: list[str]
    rows: list[list[str]]
    source: str | None = None
    section_id: str | None = None
    style_hint: str = "standard"


@dataclass(frozen=True)
class TableNormalizationResult:
    table: NormalizedTable | None
    fallback_reason: str | None = None


class TableNormalizer:
    def normalize(
        self,
        value: Any,
        *,
        title: str | None = None,
        source: str | None = None,
        section_id: str | None = None,
        style_hint: str = "standard",
        allow_single_cell: bool = False,
    ) -> NormalizedTable | None:
        result = self.normalize_with_reason(
            value,
            title=title,
            source=source,
            section_id=section_id,
            style_hint=style_hint,
            allow_single_cell=allow_single_cell,
        )
        if result.fallback_reason:
            logger.warning("DOCX table fallback: %s", result.fallback_reason)
        return result.table

    def normalize_with_reason(
        self,
        value: Any,
        *,
        title: str | None = None,
        source: str | None = None,
        section_id: str | None = None,
        style_hint: str = "standard",
        allow_single_cell: bool = False,
    ) -> TableNormalizationResult:
        resolved_style = normalize_style_hint(style_hint)
        resolved_title = clean_cell(title)
        resolved_source = clean_cell(source)
        resolved_section_id = clean_cell(section_id)
        columns: list[str] = []
        rows: list[list[str]] = []

        if isinstance(value, str):
            columns, rows = parse_pipe_table(value)
        elif isinstance(value, dict):
            resolved_title = clean_cell(value.get("title")) or resolved_title
            resolved_source = clean_cell(value.get("source")) or resolved_source
            resolved_section_id = clean_cell(value.get("section_id")) or resolved_section_id
            resolved_style = normalize_style_hint(value.get("style_hint") or resolved_style)
            columns = normalize_cells(value.get("columns") or value.get("headers"))
            rows = normalize_rows(value.get("rows") or value.get("data"))
        elif isinstance(value, list):
            rows = normalize_rows(value)
        else:
            return TableNormalizationResult(None, "unsupported table input type")

        if not columns and resolved_style in PRESET_COLUMNS:
            expected_columns = PRESET_COLUMNS[resolved_style]
            if rows and all(len(row) == len(expected_columns) for row in rows):
                columns = expected_columns

        validated = validate_table(
            columns=columns,
            rows=rows,
            allow_single_cell=allow_single_cell,
        )
        if validated.fallback_reason:
            return validated

        assert validated.table is not None
        table = NormalizedTable(
            title=resolved_title,
            columns=validated.table.columns,
            rows=validated.table.rows,
            source=resolved_source,
            section_id=resolved_section_id,
            style_hint=resolved_style,
        )
        return TableNormalizationResult(table)


class DOCXTableRenderer:
    def __init__(self, normalizer: TableNormalizer | None = None) -> None:
        self.normalizer = normalizer or TableNormalizer()

    def add_table(self, document: Any, table_data: NormalizedTable | dict[str, Any]) -> Any:
        normalized = (
            table_data
            if isinstance(table_data, NormalizedTable)
            else self.normalizer.normalize(table_data)
        )
        if normalized is None:
            return None

        if normalized.title:
            title_paragraph = document.add_paragraph()
            title_paragraph.paragraph_format.space_before = Pt(8)
            title_paragraph.paragraph_format.space_after = Pt(4)
            title_paragraph.paragraph_format.keep_with_next = True
            title_paragraph.add_run(normalized.title).bold = True

        table = document.add_table(rows=1, cols=len(normalized.columns))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table)
        set_header_repeat(table.rows[0])

        column_widths = column_widths_for(normalized)
        fill = HEADER_FILL_BY_STYLE.get(normalized.style_hint, HEADER_FILL_BY_STYLE["standard"])
        header_text_color = HEADER_TEXT_BY_STYLE.get(
            normalized.style_hint,
            HEADER_TEXT_BY_STYLE["standard"],
        )
        for index, column in enumerate(normalized.columns):
            cell = table.rows[0].cells[index]
            format_cell(
                cell,
                column,
                bold=True,
                fill=fill,
                color=header_text_color,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                width=column_widths[index],
            )

        for source_row in normalized.rows:
            row = table.add_row()
            for index, value in enumerate(source_row):
                format_cell(
                    row.cells[index],
                    value,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    width=column_widths[index],
                )

        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        return table


def normalize_style_hint(value: Any) -> str:
    if not isinstance(value, str):
        return "standard"

    normalized = value.strip().lower()
    return normalized if normalized in VALID_STYLE_HINTS else "standard"


def clean_cell(value: Any) -> str:
    if value is None:
        return ""

    return " ".join(str(value).replace("\r", "\n").split())


def normalize_cells(values: Any) -> list[str]:
    if not isinstance(values, list):
        return []

    return [clean_cell(value) for value in values]


def normalize_rows(rows: Any) -> list[list[str]]:
    if not isinstance(rows, list):
        return []

    normalized_rows: list[list[str]] = []
    for row in rows:
        if isinstance(row, dict):
            normalized_row = normalize_cells(row.get("values"))
        elif isinstance(row, list):
            normalized_row = normalize_cells(row)
        else:
            normalized_row = []

        if any(cell for cell in normalized_row):
            normalized_rows.append(normalized_row)

    return normalized_rows


def is_markdown_separator(cells: list[str]) -> bool:
    if not cells:
        return False

    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def parse_pipe_line(line: str) -> list[str]:
    cleaned = line.strip()
    if cleaned.startswith("|"):
        cleaned = cleaned[1:]
    if cleaned.endswith("|"):
        cleaned = cleaned[:-1]

    return [clean_cell(cell) for cell in cleaned.split("|")]


def parse_pipe_table(value: str) -> tuple[list[str], list[list[str]]]:
    lines = [
        line.strip()
        for line in value.splitlines()
        if line.strip() and not re.fullmatch(r"\[?Table\s+\d+[:\]]?.*", line.strip(), re.I)
    ]
    header_index = next((index for index, line in enumerate(lines) if "|" in line), None)
    if header_index is None:
        return [], []

    columns = parse_pipe_line(lines[header_index])
    data_lines = lines[header_index + 1 :]
    if data_lines and "|" in data_lines[0] and is_markdown_separator(parse_pipe_line(data_lines[0])):
        data_lines = data_lines[1:]

    rows = parse_logical_pipe_rows(data_lines, len(columns))
    return columns, rows


def parse_logical_pipe_rows(lines: list[str], column_count: int) -> list[list[str]]:
    if column_count <= 0:
        return []

    rows: list[list[str]] = []
    current_block: list[str] = []

    for line in lines:
        current_block.append(line)
        if count_row_delimiters(current_block) >= column_count - 1:
            parsed_row = parse_logical_pipe_row(current_block, column_count)
            if parsed_row:
                rows.append(parsed_row)
            current_block = []

    if current_block:
        parsed_row = parse_logical_pipe_row(current_block, column_count)
        if parsed_row:
            rows.append(parsed_row)

    return rows


def count_row_delimiters(lines: list[str]) -> int:
    text = "\n".join(lines).strip()
    if text.startswith("|") and text.endswith("|"):
        return max(text.count("|") - 2, 0)

    return text.count("|")


def parse_logical_pipe_row(lines: list[str], column_count: int) -> list[str]:
    text = "\n".join(lines).strip()
    if not text:
        return []

    if text.startswith("|") and text.endswith("|"):
        text = text[1:-1].strip()

    cells = text.split("|", maxsplit=column_count - 1)
    return [clean_cell(cell) for cell in cells]


def canonical_row(row: list[str]) -> tuple[str, ...]:
    return tuple(clean_cell(cell).lower() for cell in row)


def validate_table(
    *,
    columns: list[str],
    rows: list[list[str]],
    allow_single_cell: bool,
) -> TableNormalizationResult:
    normalized_columns = normalize_cells(columns)
    cleaned_columns = normalized_columns if any(normalized_columns) else []
    cleaned_rows = [row for row in normalize_rows(rows) if any(cell for cell in row)]
    uses_generated_columns = False

    if not cleaned_columns:
        if len(cleaned_rows) < 2:
            return TableNormalizationResult(None, "table has no reliable header row")
        candidate_header = cleaned_rows[0]
        if count_non_empty_cells(candidate_header) >= 2:
            cleaned_columns = candidate_header
            cleaned_rows = cleaned_rows[1:]
        else:
            max_column_count = max(len(row) for row in cleaned_rows)
            if max_column_count < 2 and len(cleaned_rows) <= 1 and not allow_single_cell:
                return TableNormalizationResult(None, "single-cell table is not intended")
            if max_column_count == 1 and count_non_empty_cells(candidate_header) == 1:
                cleaned_columns = candidate_header
                cleaned_rows = cleaned_rows[1:]
            else:
                cleaned_columns = [f"Column {index}" for index in range(1, max_column_count + 1)]
                uses_generated_columns = True

    column_count = len(cleaned_columns)
    if column_count == 0:
        return TableNormalizationResult(None, "table has no columns")

    if column_count == 1 and len(cleaned_rows) <= 1 and not allow_single_cell:
        return TableNormalizationResult(None, "single-cell table is not intended")

    while cleaned_rows and canonical_row(cleaned_rows[0]) == canonical_row(cleaned_columns):
        cleaned_rows.pop(0)

    if not cleaned_rows:
        return TableNormalizationResult(None, "table has no data rows")

    fixed_rows: list[list[str]] = []
    for row in cleaned_rows:
        if uses_generated_columns and len(row) < column_count:
            row = [*row, *([""] * (column_count - len(row)))]
        if len(row) != column_count:
            return TableNormalizationResult(None, "row does not match column count")
        fixed_rows.append(row)

    return TableNormalizationResult(
        NormalizedTable(title=None, columns=cleaned_columns, rows=fixed_rows)
    )


def count_non_empty_cells(row: list[str]) -> int:
    return sum(1 for cell in row if clean_cell(cell))


def get_or_add_child(parent: Any, tag: str) -> Any:
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = get_or_add_child(tbl_pr, "w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = get_or_add_child(borders, f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_cell_borders(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = get_or_add_child(tc_pr, "w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        border = get_or_add_child(borders, f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = get_or_add_child(tc_pr, "w:tcMar")
    for margin_name, width in {
        "top": "80",
        "left": "100",
        "bottom": "80",
        "right": "100",
    }.items():
        margin = get_or_add_child(margins, f"w:{margin_name}")
        margin.set(qn("w:w"), width)
        margin.set(qn("w:type"), "dxa")


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = get_or_add_child(tc_pr, "w:shd")
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell: Any, width: Any) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add_child(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_header_repeat(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_cell(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    fill: str | None = None,
    color: RGBColor | None = None,
    alignment: Any = WD_ALIGN_PARAGRAPH.LEFT,
    width: Any | None = None,
) -> None:
    cell.text = clean_cell(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if width is not None:
        set_cell_width(cell, width)
    if fill:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell)

    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color


def column_widths_for(table: NormalizedTable) -> list[Any]:
    column_count = len(table.columns)
    if column_count <= 0:
        return []

    if table.style_hint == "open_points" and column_count == 4:
        return [Inches(0.5), Inches(1.3), Inches(4.2), Inches(0.9)]
    if table.style_hint == "reports" and column_count == 3:
        return [Inches(1.8), Inches(4.0), Inches(1.2)]
    if table.style_hint == "activity" and column_count == 5:
        return [Inches(1.4), Inches(1.0), Inches(1.1), Inches(1.0), Inches(2.5)]
    if table.style_hint == "ricew" and column_count == 7:
        return [
            Inches(0.5),
            Inches(0.7),
            Inches(1.3),
            Inches(0.8),
            Inches(1.3),
            Inches(1.2),
            Inches(1.7),
        ]

    width = Inches(7.0 / column_count)
    return [width for _ in table.columns]
