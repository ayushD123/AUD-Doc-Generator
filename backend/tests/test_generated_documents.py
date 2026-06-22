import json
from base64 import b64decode
from collections.abc import Generator
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker

import app.services.docx_generation as docx_generation
from app.core.config import Settings
from app.db.base import Base
from app.db.session import get_db
from app.main import create_app
from app.models import (
    AUDPlan,
    AUDSectionDraft,
    EvidenceItem,
    ExtractedContent,
    GeneratedDocument,
    Job,
    OpenPoint,
    UploadedFile,
)
from app.services.file_storage import LocalStorageService, get_file_storage
from app.services.job_queue import LocalJobQueueService, get_job_queue_service
from app.workers.local_worker import process_generate_docx_job

ONE_PIXEL_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


def checkerboard_png_bytes() -> bytes:
    from PIL import Image

    image = Image.new("RGB", (8, 8), "white")
    for x in range(8):
        for y in range(8):
            if (x + y) % 2 == 0:
                image.putpixel((x, y), (0, 0, 0))

    output = BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


@pytest.fixture()
def client_session_and_storage(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> Generator[tuple[TestClient, sessionmaker, Path], None, None]:
    database_path = tmp_path / "test.db"
    database_url = f"sqlite:///{database_path.as_posix()}"
    storage_root = tmp_path / "storage"
    engine = create_engine(database_url, connect_args={"check_same_thread": False})
    testing_session_local = sessionmaker(
        bind=engine,
        autoflush=False,
        autocommit=False,
        expire_on_commit=False,
    )

    Base.metadata.create_all(bind=engine)
    monkeypatch.setattr(
        docx_generation,
        "get_file_storage",
        lambda: LocalStorageService(storage_root),
    )

    app = create_app(create_tables_on_startup=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as session:
            yield session

    def override_file_storage() -> LocalStorageService:
        return LocalStorageService(storage_root)

    app.dependency_overrides[get_db] = override_get_db
    app.dependency_overrides[get_file_storage] = override_file_storage
    app.dependency_overrides[get_job_queue_service] = LocalJobQueueService

    with TestClient(app) as test_client:
        yield test_client, testing_session_local, storage_root

    app.dependency_overrides.clear()
    Base.metadata.drop_all(bind=engine)
    engine.dispose()


def create_project(client: TestClient) -> str:
    response = client.post(
        "/projects",
        json={
            "name": "Asha Mehta",
            "customer_name": "Vision Operations",
            "module_name": "Order Management",
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def add_project_generation_inputs(session: Session, project_id: str) -> None:
    uploaded_file = UploadedFile(
        project_id=project_id,
        original_filename="order-management-fdd.docx",
        file_type="docx",
        storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
        source_role="fdd",
    )
    session.add(uploaded_file)
    session.flush()

    extracted_content = ExtractedContent(
        project_id=project_id,
        uploaded_file_id=uploaded_file.id,
        content_type="docx",
        title=uploaded_file.original_filename,
        text_content=(
            "[Heading: Order Capture]\n\n"
            "Orders are captured from validated FDD source material before fulfillment.\n\n"
            "[Heading: Fulfillment Flow]\n\n"
            "Fulfillment processing is documented separately."
        ),
        json_content=json.dumps(
            {
                "source_role": "fdd",
                "is_golden_source": True,
                "headings": [{"text": "Order Capture", "level": 1}],
            }
        ),
    )
    session.add(extracted_content)
    session.flush()

    plan_payload = {
        "sections": [
            {"title": "Cover Page", "include_in_aud": True},
            {"title": "Document Version History", "include_in_aud": True},
            {
                "title": "Order Capture",
                "include_in_aud": True,
                "source_role_basis": "fdd",
                "source_content_ids": [extracted_content.id],
            },
            {"title": "Open Points", "include_in_aud": True},
        ],
    }
    session.add(
        AUDPlan(
            project_id=project_id,
            status="draft",
            plan_json=json.dumps(plan_payload),
        )
    )
    session.add(
        OpenPoint(
            project_id=project_id,
            topic="Open Item",
            question="Confirm order approval threshold.",
            status="Open",
            source_type="llm_enhanced",
            refinement_status="refined",
        )
    )
    session.add(
        OpenPoint(
            project_id=project_id,
            topic="Closed Item",
            question="This resolved question should not appear.",
            status="Closed",
            source_type="llm_enhanced",
            refinement_status="refined",
        )
    )
    session.commit()


def get_docx_table_text(document: Document) -> str:
    return "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )


def create_template_docx(path: Path, paragraphs: list[str] | None = None) -> None:
    document = Document()
    for paragraph in paragraphs or [
        "Custom AUD Template",
        "<Customer Name>",
        "Oracle Fusion Cloud <Module Name>",
        "<Author>",
        "<Date>",
        "<Table Of Content>",
    ]:
        document.add_paragraph(paragraph)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def test_generate_docx_job_creates_file_and_generated_document_record(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)

    job_response = client.post(f"/projects/{project_id}/jobs/generate-docx")
    assert job_response.status_code == 201
    queued_job = job_response.json()
    assert queued_job["job_type"] == "generate_docx"
    assert queued_job["message"] == "DOCX generation job queued."

    with session_local() as session:
        job = session.scalar(select(Job).where(Job.id == queued_job["id"]))
        assert job is not None
        process_generate_docx_job(session, job, sleep_seconds=0)
        session.refresh(job)
        assert job.status == "completed"
        assert job.progress == 100

        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        assert generated_document.document_type == "aud_docx"
        assert generated_document.filename.endswith(".docx")
        output_path = storage_root / generated_document.storage_path
        assert output_path.exists()

    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Application Understanding Document" in document_text
    assert "Vision Operations" in document_text
    assert "Order Management" in document_text
    assert "Purpose and Scope" in document_text
    assert "Order Capture" in document_text
    assert (
        "Orders are captured from validated FDD source material before fulfillment."
        in document_text
    )
    assert (
        "Generated draft for Oracle internal review. Senior consultant review required "
        "before customer sharing."
    ) in document_text
    table_text = get_docx_table_text(document)
    assert "Version" in table_text
    assert "1.0" in table_text
    assert "Author" in table_text
    assert "Asha Mehta" in table_text
    assert "Confirm order approval threshold." in table_text
    assert "This resolved question should not appear." not in table_text

    list_response = client.get(f"/projects/{project_id}/generated-documents")
    assert list_response.status_code == 200
    listed_documents = list_response.json()
    assert len(listed_documents) == 1

    download_response = client.get(
        f"/projects/{project_id}/generated-documents/{listed_documents[0]['id']}/download"
    )
    assert download_response.status_code == 200
    assert download_response.content.startswith(b"PK")


def test_generated_section_headings_use_template_font_and_spacing(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "Order Capture"
    )
    run = heading.runs[0]

    assert run.font.name == docx_generation.SECTION_HEADING_FONT_NAME
    assert run.font.size.pt == docx_generation.SECTION_HEADING_FONT_SIZE_PT
    assert run.font.color.rgb == docx_generation.SECTION_HEADING_COLOR
    assert heading.paragraph_format.space_before.pt == (
        docx_generation.SECTION_HEADING_SPACE_BEFORE_PT
    )
    assert heading.paragraph_format.space_after.pt == (
        docx_generation.SECTION_HEADING_SPACE_AFTER_PT
    )
    assert heading.paragraph_format.keep_with_next is True


def test_uploaded_aud_template_is_used_when_present(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    caplog: pytest.LogCaptureFixture,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    template_storage_path = f"projects/{project_id}/uploads/uploaded-template.docx"
    create_template_docx(
        storage_root / template_storage_path,
        [
            "Uploaded template marker",
            "<Customer Name>",
            "Oracle Fusion Cloud <Module Name>",
            "Prepared by <Author>",
            "Generated on <Date>",
            "<Table Of Content>",
        ],
    )

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.add(
            UploadedFile(
                project_id=project_id,
                original_filename="uploaded-template.docx",
                file_type="docx",
                storage_path=template_storage_path,
                source_role="aud_template",
            )
        )
        session.commit()

        caplog.set_level("INFO", logger="app.services.template_resolver")
        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )

    assert "Uploaded template marker" in document_text
    assert "Vision Operations" in document_text
    assert "Oracle Fusion Cloud Order Management" in document_text
    assert "Prepared by Asha Mehta" in document_text
    assert "Order Capture" in document_text
    assert "<Customer Name>" not in document_text
    assert "Using uploaded AUD template: " in caplog.text
    assert template_storage_path in caplog.text


def test_default_template_is_used_when_no_uploaded_template_exists(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    default_template_path = tmp_path / "default-template.docx"
    create_template_docx(
        default_template_path,
        ["Default template marker", "<Customer Name>"],
    )

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        caplog.set_level("INFO", logger="app.services.template_resolver")
        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            settings=Settings(
                DEFAULT_AUD_TEMPLATE_PATH=str(default_template_path),
                _env_file=None,
            ),
        )
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )

    assert "Default template marker" in document_text
    assert "Vision Operations" in document_text
    assert "Using default AUD template: " in caplog.text
    assert str(default_template_path) in caplog.text


def test_missing_default_template_fails_with_clear_error(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    tmp_path: Path,
) -> None:
    client, session_local, _ = client_session_and_storage
    project_id = create_project(client)
    missing_template_path = tmp_path / "missing-template.docx"

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        with pytest.raises(
            FileNotFoundError,
            match="Default AUD template file not found",
        ):
            docx_generation.generate_docx(
                session,
                project_id,
                settings=Settings(
                    DEFAULT_AUD_TEMPLATE_PATH=str(missing_template_path),
                    _env_file=None,
                ),
            )


def test_generated_aud_cover_page_uses_project_metadata(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    tmp_path: Path,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    template_path = tmp_path / "metadata-template.docx"
    create_template_docx(
        template_path,
        [
            "<Customer Name>",
            "Oracle Fusion Cloud <Module Name>",
            "Version 1.0",
            "Author: <Author>",
            "Date: <Date>",
        ],
    )

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            settings=Settings(
                DEFAULT_AUD_TEMPLATE_PATH=str(template_path),
                _env_file=None,
            ),
        )
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )

    assert "Vision Operations" in document_text
    assert "Oracle Fusion Cloud Order Management" in document_text
    assert "Version 1.0" in document_text
    assert "Author: Asha Mehta" in document_text
    assert "Date: <Date>" not in document_text


def test_template_body_sample_process_sections_are_removed(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    tmp_path: Path,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    template_path = tmp_path / "sample-process-template.docx"
    create_template_docx(
        template_path,
        [
            "<Customer Name>",
            "Oracle Fusion Cloud <Module Name>",
            "Document Version History",
            "2. Process - <Process Name>",
            "Subprocess description - <Subprocess Name>",
            "<<Schedule Job List>>",
        ],
    )

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            settings=Settings(
                DEFAULT_AUD_TEMPLATE_PATH=str(template_path),
                _env_file=None,
            ),
        )
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )

    assert "2. Process -" not in document_text
    assert "Subprocess description -" not in document_text
    assert "<<Schedule Job List>>" not in document_text
    assert "<Process Name>" not in document_text


def test_unsupported_roles_and_documents_referred_are_omitted(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    tmp_path: Path,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    template_path = tmp_path / "clean-template.docx"
    create_template_docx(template_path, ["<Customer Name>", "Document Version History"])

    with session_local() as session:
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "ai_enhanced_plan": {
                            "sections": [
                                {
                                    "title": "Roles and Functions",
                                    "include_in_aud": True,
                                    "source_role_basis": "unknown",
                                    "source_content_ids": [],
                                },
                                {
                                    "title": "Documents Referred",
                                    "include_in_aud": True,
                                    "source_role_basis": "unknown",
                                    "source_content_ids": [],
                                },
                            ]
                        },
                        "sections": [
                            {"title": "Documents Referred", "include_in_aud": True}
                        ],
                    }
                ),
            )
        )
        session.commit()
        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            settings=Settings(
                DEFAULT_AUD_TEMPLATE_PATH=str(template_path),
                _env_file=None,
            ),
        )
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )

    assert "Roles and Functions" not in document_text
    assert "Documents Referred" not in document_text


def test_generated_aud_has_no_raw_template_placeholders(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    tmp_path: Path,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    template_path = tmp_path / "placeholder-template.docx"
    create_template_docx(
        template_path,
        [
            "<Customer Name>",
            "Oracle Fusion Cloud <Module Name>",
            "Document Version History",
            "<Provide the applicable content for this section.>",
            "<Placeholder text>",
            "<<Unsresolved or missing information to be listed here>>",
        ],
    )

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            settings=Settings(
                DEFAULT_AUD_TEMPLATE_PATH=str(template_path),
                _env_file=None,
            ),
        )
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = get_docx_table_text(document)
    combined_text = f"{document_text}\n{table_text}"

    assert "<Provide the applicable content for this section.>" not in combined_text
    assert "<Placeholder text>" not in combined_text
    assert "<<Unsresolved or missing information to be listed here>>" not in combined_text
    assert "<Customer Name>" not in combined_text
    assert "<Module Name>" not in combined_text


def test_docx_open_points_prefers_llm_enhanced_and_excludes_raw(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.add(
            OpenPoint(
                project_id=project_id,
                topic="Raw",
                question="Raw extracted item should not appear.",
                status="Open",
                source_type="raw_extracted",
                refinement_status="pending",
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(session, project_id)
        metadata = json.loads(generated_document.metadata_json or "{}")
        output_path = storage_root / generated_document.storage_path

    table_text = get_docx_table_text(Document(output_path))
    assert "Confirm order approval threshold." in table_text
    assert "Raw extracted item should not appear." not in table_text
    assert metadata["open_points_fallback"] is False


def test_docx_open_points_raw_fallback_is_not_rendered(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    caplog: pytest.LogCaptureFixture,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.query(OpenPoint).filter(OpenPoint.project_id == project_id).delete()
        session.add(
            OpenPoint(
                project_id=project_id,
                topic="Raw",
                question="Confirm warehouse calendar ownership.",
                status="Open",
                source_type="raw_extracted",
                refinement_status="pending",
            )
        )
        session.commit()

        pending_document = docx_generation.generate_docx(session, project_id)
        pending_path = storage_root / pending_document.storage_path
        pending_table_text = get_docx_table_text(Document(pending_path))

        raw_point = session.scalar(
            select(OpenPoint).where(OpenPoint.project_id == project_id)
        )
        assert raw_point is not None
        raw_point.refinement_status = "failed"
        session.commit()

        caplog.set_level("WARNING", logger="app.services.docx_generation")
        fallback_document = docx_generation.generate_docx(session, project_id)
        fallback_metadata = json.loads(fallback_document.metadata_json or "{}")
        fallback_path = storage_root / fallback_document.storage_path
        fallback_table_text = get_docx_table_text(Document(fallback_path))

        disabled_document = docx_generation.generate_docx(
            session,
            project_id,
            settings=Settings(ALLOW_RAW_OPEN_POINTS_FALLBACK=False, _env_file=None),
        )
        disabled_path = storage_root / disabled_document.storage_path

    disabled_table_text = get_docx_table_text(Document(disabled_path))

    assert "Confirm warehouse calendar ownership." not in pending_table_text
    assert "Confirm warehouse calendar ownership." not in fallback_table_text
    assert "Confirm warehouse calendar ownership." not in disabled_table_text
    assert fallback_metadata["open_points_fallback"] is False
    assert (
        "LLM Open Points enhancement failed; falling back to raw Open Points"
        not in caplog.text
    )


def test_open_points_table_uses_required_columns_and_table_grid(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    open_points_table = document.tables[-1]
    headers = [cell.text for cell in open_points_table.rows[0].cells]

    assert headers == ["ID", "Topic", "Question", "Status"]
    assert len(open_points_table.columns) == 4
    assert open_points_table.style.name == "Table Grid"


def test_accepted_ai_draft_is_preferred_over_rule_based_content(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        table_evidence = EvidenceItem(
            project_id=project_id,
            evidence_type="table",
            source_role="fdd",
            title="Order Capture table",
            text="Field | Value\nCapture Mode | Manual",
            json_data=json.dumps(
                {
                    "table": {
                        "rows": [
                            ["Field", "Value"],
                            ["Capture Mode", "Manual"],
                        ]
                    }
                }
            ),
            priority=100,
            confidence="high",
        )
        session.add(table_evidence)
        session.flush()
        session.add(
            AUDSectionDraft(
                project_id=project_id,
                section_id="section-order-capture",
                title="Order Capture",
                draft_text="AI accepted order capture narrative is used for the AUD.",
                draft_json=json.dumps(
                    {
                        "section_id": "section-order-capture",
                        "title": "Order Capture",
                        "included_tables": [
                            {"evidence_item_id": table_evidence.id}
                        ],
                        "included_images": [],
                    }
                ),
                confidence="high",
                review_status="accepted",
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "AI accepted order capture narrative is used for the AUD." in document_text
    assert (
        "Orders are captured from validated FDD source material before fulfillment."
        not in document_text
    )
    table_text = "\n".join(
        cell.text
        for table in Document(output_path).tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Capture Mode" in table_text
    assert "Manual" in table_text


def test_structured_source_table_does_not_render_as_plain_text_dump(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    caplog: pytest.LogCaptureFixture,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        uploaded_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        session.add(uploaded_file)
        session.flush()
        extracted_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=uploaded_file.id,
            content_type="docx",
            title=uploaded_file.original_filename,
            text_content=(
                "[Heading: Order Capture]\n\n"
                "Order capture setup is confirmed.\n\n"
                "[Table 1]\n"
                "Field | Value\n"
                "Capture Mode | Manual"
            ),
            json_content=json.dumps(
                {
                    "source_role": "fdd",
                    "is_golden_source": True,
                    "headings": [{"text": "Order Capture", "level": 1}],
                    "tables": [
                        {
                            "index": 1,
                            "section_title": "Order Capture",
                            "rows": [
                                ["Field", "Value"],
                                ["Capture Mode", "Manual"],
                            ],
                        }
                    ],
                }
            ),
        )
        session.add(extracted_content)
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Order Capture",
                                "include_in_aud": True,
                                "source_role_basis": "fdd",
                                "source_content_ids": [extracted_content.id],
                            }
                        ]
                    }
                ),
            )
        )
        session.commit()

        caplog.set_level("WARNING", logger="app.services.docx_generation")
        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = get_docx_table_text(document)

    assert "[Table 1]" not in paragraph_text
    assert "Field | Value" not in paragraph_text
    assert "Capture Mode | Manual" not in paragraph_text
    assert "Field" in table_text
    assert "Value" in table_text
    assert "Capture Mode" in table_text
    assert "Manual" in table_text
    assert "DOCX table fallback" not in caplog.text


def test_long_section_content_is_retained_without_summary_label(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        uploaded_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        session.add(uploaded_file)
        session.flush()
        section_paragraphs = [
            "[Heading: Order Details]",
            "Order capture follows the validated process definition.",
            "Order validation checks account, site, and order type before booking.",
            "Pricing validation confirms list price, discounts, and manual adjustments.",
            "Fulfillment validation confirms orchestration process assignment.",
            "Billing validation confirms invoice trigger and billing cycle.",
            "Return validation confirms RMA type and credit behavior.",
        ]
        extracted_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=uploaded_file.id,
            content_type="docx",
            title=uploaded_file.original_filename,
            text_content="\n\n".join(section_paragraphs),
            json_content=json.dumps(
                {
                    "source_role": "fdd",
                    "is_golden_source": True,
                    "headings": [{"text": "Order Details", "level": 1}],
                }
            ),
        )
        session.add(extracted_content)
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Order Details",
                                "include_in_aud": True,
                                "source_role_basis": "fdd",
                                "source_content_ids": [extracted_content.id],
                            }
                        ]
                    }
                ),
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Additional details available in source document." not in paragraph_text
    assert "Additional source details summarized:" not in paragraph_text
    assert "Billing validation confirms invoice trigger and billing cycle." in paragraph_text
    assert "Return validation confirms RMA type and credit behavior." in paragraph_text


def test_process_flow_uses_step_headings_and_bullets(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        uploaded_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        session.add(uploaded_file)
        session.flush()
        extracted_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=uploaded_file.id,
            content_type="docx",
            title=uploaded_file.original_filename,
            text_content="\n\n".join(
                [
                    "[Heading: Process - Export Order Process (Customer Care)]",
                    "Process Flow",
                    "The Export Order Process for Customer Care follows these sequential steps:",
                    "Step 1: Create Quotation",
                    "Create a Sales Order with Order Type = Quotation",
                    "Add items, freight charges, and additional information",
                    "Step 2: Create Standard Export Order",
                    "Create a Standard Export order by copying from the Quote order",
                ]
            ),
            json_content=json.dumps(
                {
                    "source_role": "fdd",
                    "is_golden_source": True,
                    "headings": [
                        {
                            "text": "Process - Export Order Process (Customer Care)",
                            "level": 1,
                        }
                    ],
                }
            ),
        )
        session.add(extracted_content)
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Process - Export Order Process (Customer Care)",
                                "include_in_aud": True,
                                "source_role_basis": "fdd",
                                "source_content_ids": [extracted_content.id],
                            }
                        ]
                    }
                ),
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    paragraphs_by_text = {paragraph.text: paragraph for paragraph in document.paragraphs}
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    bullet_texts = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "List Bullet"
        or paragraph._p.pPr is not None
        and paragraph._p.pPr.numPr is not None
    ]

    assert "Attributes" not in paragraph_text
    assert paragraphs_by_text["Process Flow"].runs[0].bold is True
    assert paragraphs_by_text["Step 1: Create Quotation"].runs[0].bold is True
    assert "Create a Sales Order with Order Type = Quotation" in bullet_texts
    assert "Add items, freight charges, and additional information" in bullet_texts


def test_draft_section_requires_include_draft_sections_option(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.add(
            AUDSectionDraft(
                project_id=project_id,
                section_id="section-order-capture",
                title="Order Capture",
                draft_text="Unreviewed AI draft should be gated.",
                draft_json=json.dumps({"included_tables": [], "included_images": []}),
                confidence="medium",
                review_status="draft",
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            options=docx_generation.DocxGenerationOptions(
                use_ai_drafts=True,
                include_draft_sections=False,
            ),
        )
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "Unreviewed AI draft should be gated." not in document_text
    assert (
        "Orders are captured from validated FDD source material before fulfillment."
        in document_text
    )


def test_generate_docx_worker_applies_job_request_options(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.add(
            AUDSectionDraft(
                project_id=project_id,
                section_id="section-order-capture",
                title="Order Capture",
                draft_text="Worker should not use this draft when option is false.",
                draft_json=json.dumps({"included_tables": [], "included_images": []}),
                confidence="medium",
                review_status="draft",
            )
        )
        job = Job(
            project_id=project_id,
            job_type="generate_docx",
            message=json.dumps(
                {
                    "status_message": "DOCX generation job queued.",
                    "options": {
                        "use_ai_drafts": True,
                        "include_draft_sections": False,
                        "include_images": True,
                        "include_open_points": True,
                    },
                }
            ),
        )
        session.add(job)
        session.commit()

        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "Worker should not use this draft when option is false." not in document_text
    assert (
        "Orders are captured from validated FDD source material before fulfillment."
        in document_text
    )


def test_omitted_ai_draft_excludes_section(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.add(
            AUDSectionDraft(
                project_id=project_id,
                section_id="section-order-capture",
                title="Order Capture",
                draft_text="This omitted draft should not appear.",
                draft_json=json.dumps({}),
                confidence="low",
                review_status="omitted",
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(session, project_id)
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "Order Capture" not in document_text
    assert "This omitted draft should not appear." not in document_text
    assert (
        "Orders are captured from validated FDD source material before fulfillment."
        not in document_text
    )


def test_ai_draft_selected_image_is_included_when_enabled(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_010_image_001.png"
    )
    image_path = storage_root / image_storage_path
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image_path.write_bytes(b64decode(ONE_PIXEL_PNG))

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        evidence_item = EvidenceItem(
            project_id=project_id,
            evidence_type="image_reference",
            source_role="kt_ppt",
            title="Order Capture selected image",
            text=image_storage_path,
            json_data=json.dumps(
                {"slide_number": 10, "image_path": image_storage_path}
            ),
            priority=70,
            confidence="high",
        )
        session.add(evidence_item)
        session.flush()
        session.add(
            AUDSectionDraft(
                project_id=project_id,
                section_id="section-order-capture",
                title="Order Capture",
                draft_text="AI accepted section with a selected image.",
                draft_json=json.dumps(
                    {
                        "included_images": [
                            {"evidence_item_id": evidence_item.id}
                        ],
                        "included_tables": [],
                    }
                ),
                confidence="high",
                review_status="accepted",
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            options=docx_generation.DocxGenerationOptions(include_images=True),
        )
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.inline_shapes) >= 1
    assert "Source image from slide 10: Order Capture selected image" in document_text


def test_docx_generator_receives_only_deduplicated_images(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    first_image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_010_image_001.png"
    )
    duplicate_image_storage_path = (
        f"projects/{project_id}/extracted_images/docx-file/image_001.png"
    )
    first_image_path = storage_root / first_image_storage_path
    duplicate_image_path = storage_root / duplicate_image_storage_path
    first_image_path.parent.mkdir(parents=True, exist_ok=True)
    duplicate_image_path.parent.mkdir(parents=True, exist_ok=True)
    first_image_path.write_bytes(b64decode(ONE_PIXEL_PNG))
    duplicate_image_path.write_bytes(b64decode(ONE_PIXEL_PNG))
    captured_image_batches: list[list[dict[str, object]]] = []
    original_add_section_images = docx_generation.add_section_images

    def capture_add_section_images(*args, **kwargs):
        images = args[1] if len(args) > 1 else kwargs.get("images", [])
        captured_image_batches.append(list(images))
        return original_add_section_images(*args, **kwargs)

    monkeypatch.setattr(
        docx_generation,
        "add_section_images",
        capture_add_section_images,
    )

    with session_local() as session:
        add_project_generation_inputs(session, project_id)
        session.add(
            AUDSectionDraft(
                project_id=project_id,
                section_id="section-order-capture",
                title="Order Capture",
                draft_text="AI accepted section with duplicate selected images.",
                draft_json=json.dumps(
                    {
                        "included_images": [
                            {
                                "image_id": "selected-ppt",
                                "storage_path": first_image_storage_path,
                                "source_uploaded_file_id": "ppt-file",
                                "source_type": "ppt_slide",
                                "slide_number": 10,
                                "caption": "Selected process screenshot",
                            },
                            {
                                "image_id": "selected-docx-copy",
                                "storage_path": duplicate_image_storage_path,
                                "source_uploaded_file_id": "docx-file",
                                "source_type": "docx_image",
                                "caption": "Selected process screenshot",
                            },
                        ],
                        "included_tables": [],
                    }
                ),
                confidence="high",
                review_status="accepted",
            )
        )
        session.commit()

        generated_document = docx_generation.generate_docx(
            session,
            project_id,
            options=docx_generation.DocxGenerationOptions(include_images=True),
        )
        metadata = json.loads(generated_document.metadata_json or "{}")

    non_empty_batches = [batch for batch in captured_image_batches if batch]
    assert len(non_empty_batches) == 1
    assert len(non_empty_batches[0]) == 1
    assert metadata["image_deduplication"]["candidate_count"] == 2
    assert metadata["image_deduplication"]["duplicates_removed_count"] == 1
    assert metadata["image_deduplication"]["retained_image_ids"] == ["selected-ppt"]


def test_missing_section_content_is_omitted_without_raw_placeholder(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        uploaded_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        session.add(uploaded_file)
        session.flush()

        extracted_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=uploaded_file.id,
            content_type="docx",
            title=uploaded_file.original_filename,
            text_content="[Heading: Available Section]\n\nAvailable FDD content.",
            json_content=json.dumps({"source_role": "fdd"}),
        )
        session.add(extracted_content)
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Missing Section",
                                "include_in_aud": True,
                                "source_role_basis": "fdd",
                                "source_content_ids": [extracted_content.id],
                            }
                        ]
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "Missing Section" not in document_text
    assert "<Content not available in provided source material>" not in document_text


def test_fdd_content_has_priority_when_mapped_with_ppt(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        fdd_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        ppt_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-kt.pptx",
            file_type="pptx",
            storage_path=f"projects/{project_id}/uploads/order-management-kt.pptx",
            source_role="kt_ppt",
        )
        session.add_all([fdd_file, ppt_file])
        session.flush()

        fdd_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=fdd_file.id,
            content_type="docx",
            title=fdd_file.original_filename,
            text_content=(
                "[Heading: Order Capture]\n\n"
                "FDD-approved order capture content wins for this section."
            ),
            json_content=json.dumps({"source_role": "fdd", "is_golden_source": True}),
        )
        ppt_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=ppt_file.id,
            content_type="pptx",
            title=ppt_file.original_filename,
            text_content="Slide 1\nTitle: Order Capture\nText:\n- PPT-only order capture content.",
            json_content=json.dumps(
                {
                    "source_role": "kt_ppt",
                    "slides": [
                        {
                            "slide_number": 1,
                            "title": "Order Capture",
                            "texts": ["PPT-only order capture content."],
                            "tables": [],
                            "notes": None,
                        }
                    ],
                }
            ),
        )
        session.add_all([fdd_content, ppt_content])
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Order Capture",
                                "include_in_aud": True,
                                "source_role_basis": "kt_ppt",
                                "source_content_ids": [ppt_content.id, fdd_content.id],
                            }
                        ]
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "FDD-approved order capture content wins for this section." in document_text
    assert "PPT-only order capture content." not in document_text


def test_docx_generation_includes_ppt_support_sections_with_fdd(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_003_image_001.png"
    )
    image_path = storage_root / image_storage_path
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image_path.write_bytes(b64decode(ONE_PIXEL_PNG))

    with session_local() as session:
        fdd_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        ppt_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-kt.pptx",
            file_type="pptx",
            storage_path=f"projects/{project_id}/uploads/order-management-kt.pptx",
            source_role="kt_ppt",
        )
        session.add_all([fdd_file, ppt_file])
        session.flush()

        session.add(
            ExtractedContent(
                project_id=project_id,
                uploaded_file_id=fdd_file.id,
                content_type="docx",
                title=fdd_file.original_filename,
                text_content=(
                    "[Heading: Order Capture]\n\n"
                    "FDD order capture content is the golden section text."
                ),
                json_content=json.dumps(
                    {
                        "source_role": "fdd",
                        "is_golden_source": True,
                        "headings": [{"text": "Order Capture", "level": 1}],
                    }
                ),
            )
        )
        session.add(
            ExtractedContent(
                project_id=project_id,
                uploaded_file_id=ppt_file.id,
                content_type="pptx",
                title=ppt_file.original_filename,
                text_content="Slide 3\nTitle: Pricing Assignments",
                json_content=json.dumps(
                    {
                        "source_role": "kt_ppt",
                        "slides": [
                            {
                                "slide_number": 3,
                                "title": "Pricing Assignments",
                                "texts": ["Pricing strategy assignment precedence."],
                                "tables": [],
                                "notes": None,
                                "image_count": 1,
                                "image_paths": [image_storage_path],
                            }
                        ],
                        "image_paths": [image_storage_path],
                        "total_image_count": 1,
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        aud_plan = session.scalar(select(AUDPlan).where(AUDPlan.project_id == project_id))
        assert generated_document is not None
        assert aud_plan is not None
        output_path = storage_root / generated_document.storage_path

    plan_payload = json.loads(aud_plan.plan_json)
    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert plan_payload["generation_basis"] == "fdd_headings_with_ppt_support"
    assert "FDD order capture content is the golden section text." in document_text
    assert "Pricing Assignments" in document_text
    assert "Pricing strategy assignment precedence." in document_text
    assert "Source image from slide 3: Pricing Assignments" in document_text
    assert len(document.inline_shapes) >= 1


def test_docx_generation_carries_enterprise_structure_content_and_image(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    image_storage_path = (
        f"projects/{project_id}/extracted_images/fdd-file/image_001.png"
    )
    image_path = storage_root / image_storage_path
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image_path.write_bytes(b64decode(ONE_PIXEL_PNG))

    with session_local() as session:
        fdd_file = UploadedFile(
            project_id=project_id,
            original_filename="order-management-fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/order-management-fdd.docx",
            source_role="fdd",
        )
        session.add(fdd_file)
        session.flush()
        session.add(
            ExtractedContent(
                project_id=project_id,
                uploaded_file_id=fdd_file.id,
                content_type="docx",
                title=fdd_file.original_filename,
                text_content=(
                    "[Heading: Introduction]\n\n"
                    "Introductory content.\n\n"
                    "[Heading: Enterprise Structure]\n\n"
                    "Business Unit: IT_BLCM_EUR_BU\n\n"
                    "[Table 1]\n"
                    "Legal Entity | Biolchim S.p.A.\n\n"
                    "[Image: "
                    f"{image_storage_path}"
                    "]\n\n"
                    "[Heading: Order Management]\n\n"
                    "Order content."
                ),
                json_content=json.dumps(
                    {
                        "source_role": "fdd",
                        "is_golden_source": True,
                        "headings": [
                            {"text": "Introduction", "level": 1},
                            {"text": "Enterprise Structure", "level": None},
                            {"text": "Order Management", "level": 1},
                        ],
                        "images": [
                            {
                                "index": 1,
                                "storage_path": image_storage_path,
                                "section_title": "Enterprise Structure",
                            }
                        ],
                        "image_paths": [image_storage_path],
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]

    assert headings.index("Introduction") < headings.index("Enterprise Structure")
    assert "Business Unit: IT_BLCM_EUR_BU" in document_text
    assert "Legal Entity | Biolchim S.p.A." in document_text
    assert "Source image from DOCX section: Enterprise Structure" in document_text
    assert len(document.inline_shapes) >= 1


def test_stale_standard_plan_is_refreshed_before_docx_generation(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)

    with session_local() as session:
        fdd_file = UploadedFile(
            project_id=project_id,
            original_filename="fdd.docx",
            file_type="docx",
            storage_path=f"projects/{project_id}/uploads/fdd.docx",
            source_role="fdd",
        )
        session.add(fdd_file)
        session.flush()

        session.add(
            ExtractedContent(
                project_id=project_id,
                uploaded_file_id=fdd_file.id,
                content_type="docx",
                title=fdd_file.original_filename,
                text_content=(
                    "[Heading: Topical Essay]\n\n"
                    "The FDD topical essay content should appear after plan refresh."
                ),
                json_content=json.dumps(
                    {
                        "source_role": "fdd",
                        "is_golden_source": True,
                        "headings": [{"text": "Topical Essay", "level": 2}],
                    }
                ),
            )
        )
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "generation_basis": "standard_sections_only",
                        "sections": [
                            {"title": "Cover Page", "include_in_aud": True},
                            {
                                "title": "Document Version History",
                                "include_in_aud": True,
                            },
                            {"title": "Table of Contents", "include_in_aud": True},
                            {"title": "Open Points", "include_in_aud": True},
                        ],
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path
        refreshed_plans = session.scalars(
            select(AUDPlan).where(AUDPlan.project_id == project_id)
        ).all()
        assert len(refreshed_plans) == 2

    document_text = "\n".join(
        paragraph.text for paragraph in Document(output_path).paragraphs
    )
    assert "Topical Essay" in document_text
    assert "The FDD topical essay content should appear after plan refresh." in document_text


def test_ppt_images_are_added_for_matching_section(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    unsupported_image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_001_image_001.wmf"
    )
    image_storage_path = f"projects/{project_id}/extracted_images/ppt-file/slide_001_image_001.png"
    low_value_image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_002_image_001.png"
    )
    mapped_image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_003_image_001.png"
    )
    unsupported_image_path = storage_root / unsupported_image_storage_path
    image_path = storage_root / image_storage_path
    low_value_image_path = storage_root / low_value_image_storage_path
    mapped_image_path = storage_root / mapped_image_storage_path
    image_path.parent.mkdir(parents=True, exist_ok=True)
    unsupported_image_path.write_bytes(b"unsupported-wmf")
    image_path.write_bytes(b64decode(ONE_PIXEL_PNG))
    low_value_image_path.write_bytes(b64decode(ONE_PIXEL_PNG))
    mapped_image_path.write_bytes(checkerboard_png_bytes())

    with session_local() as session:
        ppt_file = UploadedFile(
            project_id=project_id,
            original_filename="order-flow.pptx",
            file_type="pptx",
            storage_path=f"projects/{project_id}/uploads/order-flow.pptx",
            source_role="unknown",
        )
        session.add(ppt_file)
        session.flush()

        ppt_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=ppt_file.id,
            content_type="pptx",
            title=ppt_file.original_filename,
            text_content="Slide 1\nTitle: Fulfillment Flow",
            json_content=json.dumps(
                {
                    "source_role": "unknown",
                    "slides": [
                        {
                            "slide_number": 1,
                            "title": "Fulfillment Flow",
                            "texts": ["Reserve inventory and release shipment."],
                            "tables": [],
                            "notes": None,
                            "image_count": 2,
                            "image_paths": [
                                unsupported_image_storage_path,
                                image_storage_path,
                            ],
                        },
                        {
                            "slide_number": 2,
                            "title": "Thank You",
                            "texts": ["Thank You"],
                            "tables": [],
                            "notes": None,
                            "image_count": 1,
                            "image_paths": [low_value_image_storage_path],
                        },
                        {
                            "slide_number": 3,
                            "title": "Configuration Snapshot",
                            "texts": ["High-value screenshot with process labels."],
                            "tables": [],
                            "notes": None,
                            "image_count": 1,
                            "image_paths": [mapped_image_storage_path],
                        },
                    ],
                    "image_paths": [
                        unsupported_image_storage_path,
                        image_storage_path,
                        low_value_image_storage_path,
                        mapped_image_storage_path,
                    ],
                    "total_image_count": 4,
                }
            ),
        )
        session.add(ppt_content)
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Fulfillment Flow",
                                "include_in_aud": True,
                                "source_role_basis": "kt_ppt",
                                "source_content_ids": [ppt_content.id],
                            },
                            {
                                "title": "Exception Handling",
                                "include_in_aud": True,
                                "source_role_basis": "kt_ppt",
                                "source_content_ids": [ppt_content.id],
                                "notes": ["Source slide 3."],
                            },
                        ]
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.inline_shapes) >= 2
    assert "Source image from slide 1: Fulfillment Flow" in document_text
    assert "Source image from slide 3: Configuration Snapshot" in document_text
    assert "Source image from slide 2: Thank You" not in document_text


def test_single_word_ppt_title_does_not_overmatch_compound_section(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, session_local, storage_root = client_session_and_storage
    project_id = create_project(client)
    service_image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_001_image_001.png"
    )
    pricing_image_storage_path = (
        f"projects/{project_id}/extracted_images/ppt-file/slide_002_image_001.png"
    )
    service_image_path = storage_root / service_image_storage_path
    pricing_image_path = storage_root / pricing_image_storage_path
    service_image_path.parent.mkdir(parents=True, exist_ok=True)
    service_image_path.write_bytes(b64decode(ONE_PIXEL_PNG))
    pricing_image_path.write_bytes(b64decode(ONE_PIXEL_PNG))

    with session_local() as session:
        ppt_file = UploadedFile(
            project_id=project_id,
            original_filename="order-flow.pptx",
            file_type="pptx",
            storage_path=f"projects/{project_id}/uploads/order-flow.pptx",
            source_role="kt_ppt",
        )
        session.add(ppt_file)
        session.flush()

        ppt_content = ExtractedContent(
            project_id=project_id,
            uploaded_file_id=ppt_file.id,
            content_type="pptx",
            title=ppt_file.original_filename,
            text_content="Slide 1\nTitle: Service Mappings (OM/Pricing)",
            json_content=json.dumps(
                {
                    "source_role": "kt_ppt",
                    "slides": [
                        {
                            "slide_number": 1,
                            "title": "Service Mappings (OM/Pricing)",
                            "texts": ["Service mapping details."],
                            "tables": [],
                            "notes": None,
                            "image_count": 1,
                            "image_paths": [service_image_storage_path],
                        },
                        {
                            "slide_number": 2,
                            "title": "Pricing",
                            "texts": ["Pricing details."],
                            "tables": [],
                            "notes": None,
                            "image_count": 1,
                            "image_paths": [pricing_image_storage_path],
                        },
                    ],
                    "image_paths": [
                        service_image_storage_path,
                        pricing_image_storage_path,
                    ],
                    "total_image_count": 2,
                }
            ),
        )
        session.add(ppt_content)
        session.flush()
        session.add(
            AUDPlan(
                project_id=project_id,
                status="draft",
                plan_json=json.dumps(
                    {
                        "sections": [
                            {
                                "title": "Service Mappings (OM/Pricing)",
                                "include_in_aud": True,
                                "source_role_basis": "kt_ppt",
                                "source_content_ids": [ppt_content.id],
                            }
                        ]
                    }
                ),
            )
        )
        session.commit()

    with session_local() as session:
        job = Job(project_id=project_id, job_type="generate_docx")
        session.add(job)
        session.commit()
        process_generate_docx_job(session, job, sleep_seconds=0)
        generated_document = session.scalar(
            select(GeneratedDocument).where(GeneratedDocument.project_id == project_id)
        )
        assert generated_document is not None
        output_path = storage_root / generated_document.storage_path

    document = Document(output_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert len(document.inline_shapes) >= 1
    assert "Source image from slide 1: Service Mappings (OM/Pricing)" in document_text
    assert "Source image from slide 2: Pricing" not in document_text


def test_list_generated_documents_returns_404_for_unknown_project(
    client_session_and_storage: tuple[TestClient, sessionmaker, Path],
) -> None:
    client, _, _ = client_session_and_storage

    response = client.get("/projects/missing-project/generated-documents")

    assert response.status_code == 404
    assert response.json()["detail"] == "Project not found."
