import logging

from docx import Document

from app.services.docx_table_renderer import DOCXTableRenderer, TableNormalizer


def table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def test_markdown_table_becomes_docx_table() -> None:
    document = Document()
    table = TableNormalizer().normalize(
        """
| Column A | Column B |
| --- | --- |
| value 1 | value 2 |
"""
    )

    assert table is not None
    DOCXTableRenderer().add_table(document, table)

    assert len(document.tables) == 1
    assert table_text(document.tables[0]) == [
        ["Column A", "Column B"],
        ["value 1", "value 2"],
    ]


def test_pipe_delimited_table_becomes_docx_table() -> None:
    document = Document()
    table = TableNormalizer().normalize("Column A | Column B\nvalue 1 | value 2")

    assert table is not None
    DOCXTableRenderer().add_table(document, table)

    assert len(document.tables) == 1
    assert table_text(document.tables[0]) == [
        ["Column A", "Column B"],
        ["value 1", "value 2"],
    ]


def test_multiline_pipe_cells_and_leading_blank_cells_are_preserved() -> None:
    table = TableNormalizer().normalize(
        """Name | IF Condition Description | Process Name
Orchestration
Process
Assignment Rule | Order Type = Standard Domestic
Order Type = Standard EU
AND
Line Type = Buy
Line Type = Sample | CustomDOO_DOO_OrderFulfillmentGenericProcess_NoRSV
| Order Type = Quotation | CustomDOO_DOO_OrderFulfillmentGenericProcess_CLOSE_ONLY"""
    )

    assert table is not None
    assert table.columns == ["Name", "IF Condition Description", "Process Name"]
    assert table.rows == [
        [
            "Orchestration Process Assignment Rule",
            "Order Type = Standard Domestic Order Type = Standard EU AND Line Type = Buy Line Type = Sample",
            "CustomDOO_DOO_OrderFulfillmentGenericProcess_NoRSV",
        ],
        [
            "",
            "Order Type = Quotation",
            "CustomDOO_DOO_OrderFulfillmentGenericProcess_CLOSE_ONLY",
        ],
    ]


def test_open_points_table_has_required_columns() -> None:
    table = TableNormalizer().normalize(
        {
            "style_hint": "open_points",
            "rows": [["1", "Topic", "Question?", "Open"]],
        }
    )

    assert table is not None
    assert table.columns == ["ID", "Topic", "Question", "Status"]


def test_activity_table_has_expected_columns() -> None:
    table = TableNormalizer().normalize(
        {
            "style_hint": "activity",
            "rows": [["Capture order", "Ops", "Manual", "OM", "Create order"]],
        }
    )

    assert table is not None
    assert table.columns == [
        "Activity",
        "Owner",
        "Automatic/Manual",
        "System",
        "Description",
    ]


def test_single_column_multi_row_table_is_allowed() -> None:
    table = TableNormalizer().normalize(
        {
            "title": "Process List",
            "rows": [
                ["Process"],
                ["Customize Orchestration Process"],
                ["Order Header and Line Level EFFs"],
            ],
        }
    )

    assert table is not None
    assert table.columns == ["Process"]
    assert table.rows == [
        ["Customize Orchestration Process"],
        ["Order Header and Line Level EFFs"],
    ]


def test_blank_header_source_table_gets_generated_columns() -> None:
    table = TableNormalizer().normalize(
        {
            "title": "Line Level Details",
            "rows": [
                ["", "", ""],
                ["Line Level", "", ""],
                ["Attribute", "Service Mapping", "Algorithm"],
            ],
        }
    )

    assert table is not None
    assert table.columns == ["Column 1", "Column 2", "Column 3"]
    assert table.rows == [
        ["Line Level", "", ""],
        ["Attribute", "Service Mapping", "Algorithm"],
    ]


def test_cell_borders_are_present_in_generated_xml() -> None:
    document = Document()
    table = TableNormalizer().normalize("Column A | Column B\nvalue 1 | value 2")

    assert table is not None
    DOCXTableRenderer().add_table(document, table)

    xml = document.tables[0]._tbl.xml
    assert "w:tblBorders" in xml
    assert "w:tcBorders" in xml


def test_malformed_table_falls_back_with_warning(
    caplog,
) -> None:
    caplog.set_level(logging.WARNING, logger="app.services.docx_table_renderer")

    table = TableNormalizer().normalize("Column A | Column B\nvalue 1")

    assert table is None
    assert "DOCX table fallback" in caplog.text
